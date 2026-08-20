from pathlib import Path
from datetime import date
import json
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "LiveInBlack_Specification_Fonctionnelle.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_run(run, size=11, color="17202A", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    style_run(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    style_run(p.add_run(text))
    return p


def add_para(doc, text="", bold_prefix=None, color="17202A"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        style_run(p.add_run(bold_prefix), bold=True, color="0B2545")
        style_run(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        style_run(p.add_run(text), color=color)
    return p


def add_callout(doc, label, text, fill="E8F1F8"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(label + "  "), bold=True, color="0B2545")
    style_run(p.add_run(text), color="17202A")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(header), bold=True, color="0B2545")
    repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            style_run(p.add_run(str(value)))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def route_label(path):
    return path.replace("/page.tsx", "").replace("/route.ts", "")


def collect_routes():
    pages = sorted(str(p.relative_to(ROOT)).replace("\\", "/") for p in (ROOT / "app").rglob("page.tsx"))
    api = sorted(str(p.relative_to(ROOT)).replace("\\", "/") for p in (ROOT / "app/api").rglob("route.ts"))
    return pages, api


def add_route_catalog(doc, title, entries, prefix_filter=None, limit=None):
    doc.add_heading(title, level=2)
    filtered = [e for e in entries if not prefix_filter or prefix_filter in e]
    if limit:
        filtered = filtered[:limit]
    for entry in filtered:
        display = entry
        if display.startswith("app/"):
            display = display[4:]
        display = re.sub(r"^\([^)]*\)/", "", display)
        display = display.replace("/page.tsx", "").replace("/route.ts", "")
        if display.startswith("api/"):
            display = "/" + display
        add_bullet(doc, display)
    return len(filtered)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run("LIVEINBLACK  |  Spécification fonctionnelle"), size=9, color="6B7785")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer.add_run("Document de référence — généré le " + date.today().isoformat()), size=9, color="6B7785")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    pages, api = collect_routes()
    package = json.loads((ROOT / "package.json").read_text())

    # Cover / masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run("LIVEINBLACK"), size=13, bold=True, color="E45757")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run("Spécification fonctionnelle exhaustive"), size=27, bold=True, color="0B2545")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    style_run(p.add_run("Produit web de découverte, réservation et animation d’événements"), size=14, color="52606D")
    add_table(doc, ["Version", "Date", "Périmètre", "Référence"], [[package.get("version", "0.1.0"), date.today().isoformat(), "Web public, espaces connectés, agent et API", "Code source courant"]], [1500, 1500, 4300, 2060])
    add_callout(doc, "Objet", "Ce document décrit ce que fait l’application, qui peut le faire, par quels écrans et quelles API, ainsi que les règles métier et contrôles attendus pour l’exploitation et la recette.")

    doc.add_heading("1. Résumé exécutif", level=1)
    add_para(doc, "LIVEINBLACK est une plateforme de sorties et d’événements. Elle met en relation des clients, des organisateurs, des prestataires et des agents de contrôle autour de la découverte d’événements, la réservation de billets, la messagerie et la modération.")
    add_bullet(doc, "Découvrir : événements, organisateurs, prestataires, blog et contenus d’aide.")
    add_bullet(doc, "Convertir : inscription, connexion, achat de billets, invitations, revente et paiements.")
    add_bullet(doc, "Animer : intérêts, abonnements, amis, notifications, playlists et messagerie.")
    add_bullet(doc, "Opérer : studios organisateur/prestataire, ventes sur place, scanner, paiements et espace agent.")
    add_bullet(doc, "Protéger : rôles, contrôles d’accès, validation des webhooks, confidentialité, export et suppression de compte.")

    doc.add_heading("2. Périmètre et architecture fonctionnelle", level=1)
    add_table(doc, ["Couche", "Responsabilité", "Surface"], [
        ("Web public", "Acquisition, découverte et information", "home, events, organizers, providers, blog, aide, légal"),
        ("Web connecté", "Profil, billets, social, messagerie et espaces métier", "client, organisateur, prestataire, agent"),
        ("API Next.js", "Authentification, règles métier, données et intégrations", f"{len(api)} routes API détectées"),
        ("Données", "Persistance des utilisateurs, événements, commandes, tickets, social et modération", "MongoDB / Mongoose"),
        ("Services externes", "Paiements, email, médias, musique, analytics et déploiement", "Stripe, FedaPay, Resend, Cloudinary, Deezer/Apple, Vercel"),
    ], [1600, 3900, 3860])
    add_para(doc, "Le client mobile consomme les mêmes routes web via une base API configurable. Le contrat statique courant relie 207 appels mobiles uniques à une route et une méthode, dont 31 appels de l’espace agent.")

    doc.add_heading("3. Rôles et responsabilités", level=1)
    add_table(doc, ["Rôle", "Objectif", "Capacités principales"], [
        ("Visiteur", "Explorer sans compte", "Consulter pages publiques, événements, profils, blog, aide et légal; rechercher; ouvrir les modales publiques."),
        ("Client", "Participer à une sortie", "Acheter/recevoir des billets, gérer profil, intérêts, organisateurs suivis, amis, notifications, messages, revente et remboursements."),
        ("Organisateur", "Créer et exploiter ses événements", "Studio, événements, billets, invités, équipe, promo, playlist, statistiques, ventes et reversements."),
        ("Prestataire", "Présenter et vendre ses services", "Profil public, catalogue, médias, avis, abonnement et facturation."),
        ("Agent", "Modérer et administrer la plateforme", "Dashboard, candidatures, comptes, événements, signalements, avis, paiements, suppressions, homepage, blog et boosts."),
    ], [1500, 2500, 5360])
    add_callout(doc, "Règle d’accès", "Une action connectée doit renvoyer une erreur d’authentification si la session est absente, puis appliquer le rôle et la propriété de la ressource. L’espace agent est réservé aux comptes portant le rôle agent.", "FFF4D6")

    doc.add_heading("4. Plan du site et écrans", level=1)
    doc.add_heading("4.1 Pages publiques", level=2)
    for label in [
        "/home — accueil et découverte", "/events — annuaire et recherche d’événements", "/events/[id] — détail d’un événement", "/organizers — annuaire des organisateurs", "/organizers/[slug] — profil public organisateur", "/providers — annuaire des prestataires", "/providers/[id] — profil public prestataire", "/blog et /blog/[slug] — contenus éditoriaux", "/about — présentation", "/help — aide et fonctionnement", "/contact — contact", "/cookies, /privacy, /terms, /legal-notice — information légale", "/login, /reset-password, /verify-email, /confirmer-email — accès et compte", "/organizer-signup et /provider-signup — candidatures métier", "/payment-success, /boost-active — retours transactionnels",
    ]:
        add_bullet(doc, label)
    doc.add_heading("4.2 Pages connectées client", level=2)
    for label in [
        "/profile — synthèse du profil", "/profile/parametres — compte, confidentialité, préférences et sécurité", "/profile/billets — billets possédés", "/profile/interested-events — événements intéressés", "/profile/followed-organizers — organisateurs suivis", "/messages — conversations", "/notifications — centre de notifications", "/help — aide contextualisée", "/order/[eventId]/[ticketCode] — billet et contrôle", "/checkout — réservation et paiement", "/playlist/[eventId] — playlist participative", "/my-application — suivi de candidature", "/spaces — choix d’espace métier",
    ]:
        add_bullet(doc, label)
    doc.add_heading("4.3 Pages organisateur, prestataire et agent", level=2)
    for label in [
        "Organisateur : /organizer-studio, /my-events, /my-events/[id]/statistiques, /my-shifts, /on-site-sales/[eventId], /scanner/[eventId]",
        "Prestataire : /offer-services et profil/catalogue métier",
        "Agent : /agent, /agent/evenements, /agent/comptes, /agent/dossiers, /agent/paiements, /agent/signalements, /agent/avis, /agent/suppressions, /agent/actualite, /agent/blog",
    ]:
        add_bullet(doc, label)

    doc.add_heading("5. Parcours fonctionnels de référence", level=1)
    workflows = [
        ("Inscription et activation", ["Le visiteur choisit client, organisateur ou prestataire.", "Le compte est créé puis l’email est vérifié.", "Le profil, les préférences et le rôle actif sont accessibles après connexion.", "Les demandes métier peuvent rester en brouillon avant soumission et modération."]),
        ("Découverte et intérêt", ["Le visiteur filtre par recherche, région, catégorie, ambiance ou date.", "La page détail présente l’événement, l’organisateur, les billets, la playlist et les actions disponibles.", "Le client connecté peut ajouter/retirer son intérêt ou suivre l’organisateur."]),
        ("Achat de billet", ["Le client sélectionne un type de billet et une quantité.", "Le système crée une commande ou une réservation de places avec expiration si nécessaire.", "Le paiement passe par Stripe/FedaPay ou le parcours gratuit.", "La confirmation produit les billets, le QR/token, l’historique et les notifications."]),
        ("Billet, invitation et contrôle", ["Un billet peut être assigné, transféré, accepté/refusé, revendu ou révoqué selon son état.", "Le scanner vérifie le token et empêche les doublons.", "Les commandes et journaux de service assurent la traçabilité."]),
        ("Organisation d’événement", ["L’organisateur configure informations, places, tarifs, médias, staff, guestlist, promo et playlist.", "La publication rend l’événement découvrable selon son statut.", "Les changements, reports et annulations appliquent des contrôles de propriété et déclenchent remboursements/notifications."]),
        ("Messagerie", ["Les utilisateurs créent une conversation directe ou un groupe.", "Ils envoient textes, médias, réactions, sondages, votes, messages épinglés et favoris.", "Les membres peuvent lire, masquer, quitter, renommer, gérer les rôles et signaler un contenu."]),
        ("Modération agent", ["L’agent consulte les files d’événements, comptes, candidatures, avis, signalements, paiements et suppressions.", "Chaque action de validation, refus, suspension, remboursement ou résolution est contrôlée par rôle et journalisée.", "Les actions sensibles envoient les notifications et emails prévus."]),
    ]
    for title, steps in workflows:
        doc.add_heading(title, level=2)
        for step in steps:
            add_number(doc, step)

    doc.add_heading("6. Modules fonctionnels détaillés", level=1)
    modules = [
        ("Identité et profil", "Inscription email, connexion Auth.js, vérification, réinitialisation, changement de mot de passe, nom, email, téléphone, avatar, démographie, préférences, confidentialité, export et suppression."),
        ("Événements", "Création, édition, publication, recherche, détail, catégories, styles musicaux, ambiances, lieux, images, vidéo, date, capacité, tarifs, annulation et report."),
        ("Billetterie", "Types de billets, commandes, billets gratuits/payants, hold de place, invitations, assignation, annulation, check-in, revente, remboursement et reçus."),
        ("Social", "Intérêt événement, abonnements organisateurs, amis, demandes, notifications, présence, blocage et signalement utilisateur."),
        ("Messagerie", "Conversations directes/groupes, membres, rôles, mute, avatar, contact téléphone, messages, médias, réactions, édition, suppression, forward, star, polls, read, typing et pin."),
        ("Organisateur", "Studio, événements, statistiques, staff, guestlist, promos, playlist, ventes sur place, shifts, payouts, historique et médias."),
        ("Prestataire", "Profil public, catalogue, médias, avis, abonnement, région de facturation et gestion de l’offre."),
        ("Agent", "Dashboard, users, applications, events, reports, reviews, deletions, payments/refunds/payouts, homepage, blog et boosts."),
        ("Éditorial et aide", "Blog, actualité homepage, contact, FAQ, pages d’information, confidentialité, cookies, CGU et mentions légales."),
    ]
    add_table(doc, ["Module", "Fonctions couvertes"], modules, [2100, 7260])

    doc.add_heading("6.1 Inventaire exhaustif des fonctionnalités", level=2)
    add_para(doc, "Chaque ligne ci-dessous constitue une fonctionnalité testable. Les identifiants F-xxx servent de référence commune entre produit, développement, support et QA.")
    feature_groups = [
        ("Identité, accès et compte", [
            ("Créer un compte client", "Visiteur", "Saisir email, mot de passe et données minimales.", "Email non utilisé; le compte est créé et peut demander la vérification."),
            ("Créer une candidature organisateur", "Visiteur/Client", "Choisir le rôle organisateur et ouvrir le formulaire progressif.", "Un brouillon de candidature est enregistré."),
            ("Créer une candidature prestataire", "Visiteur/Client", "Choisir le rôle prestataire et renseigner l’offre.", "Un brouillon de candidature est enregistré."),
            ("Se connecter", "Tous les comptes", "Soumettre email et mot de passe via Auth.js.", "Session ouverte; redirection selon le contexte et le rôle actif."),
            ("Vérifier l’adresse email", "Utilisateur", "Cliquer le lien signé reçu par email.", "Email marqué vérifié; lien expiré ou déjà utilisé refusé."),
            ("Renvoyer la vérification", "Utilisateur", "Demander un nouveau lien depuis l’écran de compte.", "Un nouveau lien est envoyé avec limite anti-abus."),
            ("Réinitialiser le mot de passe", "Utilisateur", "Demander un token puis définir un nouveau mot de passe.", "Ancien mot de passe invalidé; token à usage unique."),
            ("Changer le mot de passe", "Utilisateur connecté", "Confirmer l’ancien mot de passe et saisir le nouveau.", "Mot de passe mis à jour et sessions sensibles sécurisées."),
            ("Changer le rôle actif", "Compte multi-rôles", "Sélectionner client, organisateur, prestataire ou agent.", "Les menus et permissions changent sans modifier les rôles attribués."),
            ("Modifier identité et coordonnées", "Utilisateur connecté", "Modifier nom, email, téléphone et avatar.", "Profil persistant avec validation des formats et doublons."),
            ("Gérer confidentialité et préférences", "Utilisateur connecté", "Choisir visibilité, présence, accusés de lecture, ambiance et ville.", "Préférences appliquées au profil et à la découverte."),
            ("Exporter les données", "Utilisateur connecté", "Demander une copie des données du compte.", "Export préparé selon les données autorisées."),
            ("Demander la suppression", "Utilisateur connecté", "Confirmer le mot de passe et la demande.", "Purge immédiate ou dossier agent selon le statut métier."),
        ]),
        ("Découverte et recherche", [
            ("Afficher l’accueil", "Visiteur/Utilisateur", "Charger hero, événements recommandés, actualité et appels à l’action.", "Contenu public visible même si l’utilisateur est connecté."),
            ("Lister les événements", "Visiteur/Utilisateur", "Parcourir pagination, filtres, tri, dates, région et recherche.", "Liste cohérente; état vide et erreur explicites."),
            ("Rechercher rapidement", "Visiteur/Utilisateur", "Saisir une recherche dans le header.", "Suggestions événements, organisateurs et prestataires."),
            ("Consulter le détail événement", "Visiteur/Utilisateur", "Ouvrir la page ou la modale d’un événement.", "Détails, médias, billets, organisateur, playlist et actions."),
            ("Lister les organisateurs", "Visiteur/Utilisateur", "Filtrer par région, popularité, récence et événements à venir.", "Annuaire public paginé."),
            ("Consulter un organisateur", "Visiteur/Utilisateur", "Ouvrir le profil public par slug.", "Bio, médias, événements, vérification et bouton suivre selon session."),
            ("Lister les prestataires", "Visiteur/Utilisateur", "Filtrer par catégorie, zone, prix ou recherche.", "Catalogue public paginé."),
            ("Consulter un prestataire", "Visiteur/Utilisateur", "Ouvrir le profil public.", "Présentation, catalogue, médias, avis et contact."),
            ("Lire le blog", "Visiteur/Utilisateur", "Parcourir les articles puis ouvrir un slug.", "Article public, image, contenu et navigation éditoriale."),
            ("Consulter aide et informations", "Visiteur/Utilisateur", "Ouvrir aide, à propos, contact, cookies, confidentialité, CGU et mentions.", "Informations visibles et liens fonctionnels."),
        ]),
        ("Événements et organisation", [
            ("Créer un événement", "Organisateur", "Renseigner titre, date, lieu, description, catégories et billets.", "Brouillon créé avec propriétaire contrôlé."),
            ("Modifier un événement", "Organisateur autorisé", "Modifier les champs et médias autorisés.", "Version mise à jour; contraintes de date et capacité appliquées."),
            ("Publier/dépublier", "Organisateur autorisé", "Changer le statut de visibilité.", "L’événement entre ou sort de la découverte publique."),
            ("Gérer types et quotas de billets", "Organisateur", "Créer tarifs, quantités, périodes et règles.", "Disponibilité calculée sans dépasser la capacité."),
            ("Gérer équipe et staff", "Organisateur", "Inviter, modifier rôle, retirer membres et droits.", "Accès staff limité aux événements affectés."),
            ("Gérer guestlist", "Organisateur/Staff", "Ajouter invités, émettre, retirer ou valider des entrées.", "Billets/invitations traçables."),
            ("Créer codes promo", "Organisateur", "Définir code, remise, période, limite et billets ciblés.", "Prix recalculé et code anti-abus."),
            ("Gérer playlist", "Organisateur/DJ", "Rechercher des titres, ajouter, liker, modérer et lancer le titre courant.", "Playlist événement persistante et contrôlée."),
            ("Reporter un événement", "Organisateur autorisé", "Saisir une nouvelle date et justifier le report.", "Date validée; clients notifiés selon impact."),
            ("Annuler un événement", "Organisateur/Agent", "Confirmer l’annulation.", "Ventes bloquées; remboursements et notifications déclenchés."),
            ("Consulter statistiques", "Organisateur", "Ouvrir ventes, capacité, clics, vues et tendances.", "Indicateurs agrégés selon propriétaire."),
            ("Vendre sur place", "Staff vendeur", "Choisir événement, billet et mode de règlement.", "Commande et tickets émis; règlement journalisé."),
            ("Scanner un billet", "Staff autorisé", "Scanner ou saisir le token.", "Billet accepté une fois ou motif de refus explicite."),
        ]),
        ("Billetterie, commandes et paiements", [
            ("Réserver un billet gratuit", "Client/Visiteur", "Choisir un billet à prix nul et confirmer.", "Commande matérialisée et ticket émis sans paiement."),
            ("Payer une commande", "Client", "Choisir Stripe ou FedaPay selon disponibilité.", "Checkout créé; commande confirmée après webhook/retour validé."),
            ("Maintenir une place", "Client", "Créer un seat hold avec expiration.", "Place réservée temporairement puis libérée si expiration."),
            ("Consulter ses billets", "Client", "Ouvrir l’espace billets.", "Billets filtrés par utilisateur, événement et état."),
            ("Partager/assigner un billet", "Client", "Saisir le destinataire et confirmer.", "Invitation sortante créée; destinataire accepte ou refuse."),
            ("Quitter un billet assigné", "Client", "Demander à quitter le billet reçu.", "Billet rendu disponible selon les règles de l’événement."),
            ("Revendre un billet", "Client", "Créer une annonce avec prix et conditions.", "Annonce active; acheteur peut lancer le checkout de revente."),
            ("Acheter une revente", "Client", "Ouvrir une annonce et payer.", "Transfert sécurisé et annonce clôturée."),
            ("Demander un remboursement", "Client", "Ouvrir la commande et envoyer le motif.", "Demande idempotente; traitement selon statut de l’événement."),
            ("Rembourser/compléter côté agent", "Agent", "Ouvrir la file des remboursements et marquer le traitement.", "Alerte résolue et historique conservé."),
            ("Demander un reversement", "Organisateur", "Renseigner moyen et montant éligible.", "Payout request créé avec contrôle de solde."),
            ("Gérer l’abonnement prestataire", "Prestataire", "Souscrire, consulter ou annuler l’abonnement.", "État d’abonnement et facturation synchronisés."),
        ]),
        ("Social, notifications et messagerie", [
            ("Ajouter un intérêt événement", "Client connecté", "Cliquer intérêt depuis liste ou détail.", "Intérêt créé et visible dans le profil."),
            ("Suivre un organisateur", "Client connecté", "Cliquer suivre depuis annuaire ou profil.", "Follow créé; compteur et préférences alertes mis à jour."),
            ("Gérer alertes d’un organisateur", "Client connecté", "Activer/désactiver types d’alertes.", "Préférences de notification persistantes."),
            ("Envoyer une demande d’ami", "Utilisateur connecté", "Sélectionner un utilisateur et envoyer.", "Demande pending avec notification."),
            ("Accepter/refuser/annuler une demande", "Utilisateur connecté", "Agir depuis demandes ou notification.", "État de relation recalculé."),
            ("Bloquer/débloquer un utilisateur", "Utilisateur connecté", "Choisir bloquer depuis profil ou conversation.", "Accès social et messagerie filtrés."),
            ("Créer une conversation directe", "Utilisateur connecté", "Choisir un contact autorisé.", "Conversation réutilisée si elle existe."),
            ("Créer un groupe", "Utilisateur connecté", "Choisir membres, nom et avatar.", "Groupe créé avec admin initial."),
            ("Envoyer et gérer un message", "Membre", "Envoyer texte/média, éditer, supprimer ou transférer.", "Message visible selon droits et historique."),
            ("Réagir, favoriser et voter", "Membre", "Réagir à un message, le mettre en favori ou voter à un sondage.", "État personnel et compteurs synchronisés."),
            ("Gérer une conversation", "Membre/Admin", "Lire, mute, épingler, renommer, quitter, gérer membres et rôles.", "Paramètres appliqués au groupe ou au membre."),
            ("Consulter et traiter notifications", "Utilisateur connecté", "Lister, ouvrir, marquer lu ou tout marquer lu.", "Compteur et statut de lecture cohérents."),
        ]),
        ("Prestataire", [
            ("Compléter son profil public", "Prestataire", "Renseigner présentation, zones, catégories et médias.", "Profil publié selon statut de candidature."),
            ("Gérer le catalogue", "Prestataire", "Créer, modifier, archiver une offre et ses prix.", "Catalogue public à jour."),
            ("Gérer médias d’une offre", "Prestataire", "Ajouter, ordonner et supprimer médias signés.", "Médias privés/publics selon règle."),
            ("Consulter ses avis", "Prestataire", "Ouvrir avis reçus et réponses.", "Avis paginés et moyenne recalculée."),
            ("Répondre à un avis", "Prestataire", "Saisir une réponse sur un avis autorisé.", "Réponse publiée ou validation d’erreur explicite."),
            ("Choisir région de facturation", "Prestataire", "Sélectionner la région supportée.", "Région utilisée pour la facturation et les moyens de paiement."),
        ]),
        ("Agent et administration", [
            ("Consulter dashboard agent", "Agent", "Ouvrir les indicateurs et files en attente.", "Stats globales et compteurs chargés."),
            ("Rechercher un compte", "Agent", "Filtrer, ouvrir le détail et modifier champs autorisés.", "Compte mis à jour avec journalisation."),
            ("Vérifier/suspendre un compte", "Agent", "Désactiver, vérifier email ou envoyer reset.", "État du compte et notification actualisés."),
            ("Modérer une candidature", "Agent", "Ouvrir dossier, noter, approuver ou refuser.", "Statut métier et emails mis à jour."),
            ("Modérer un événement", "Agent", "Lister, consulter et annuler un événement problématique.", "Annulation contrôlée et remboursements traités."),
            ("Traiter un signalement", "Agent", "Lire rapport et marquer handled avec note.", "Rapport résolu et action conservée."),
            ("Modérer un avis", "Agent", "Masquer, restaurer ou traiter un avis signalé.", "Visibilité recalculée."),
            ("Traiter une suppression", "Agent", "Approuver ou rejeter une demande de compte.", "Purge/anonymisation déclenchée selon décision."),
            ("Traiter paiements et payouts", "Agent", "Consulter alertes, refunds, payouts et settle.", "État financier mis à jour et traçable."),
            ("Gérer homepage et blog", "Agent", "Modifier bandeau éditorial, articles et publication.", "Contenu public mis à jour avec contrôle de statut."),
            ("Consulter les boosts", "Agent", "Voir campagnes actives et créneaux.", "État et disponibilité des boosts visibles."),
        ]),
        ("Système, intégrations et conformité", [
            ("Recevoir un webhook Stripe", "Stripe", "Vérifier signature et traiter l’événement.", "Réponse 2xx après traitement idempotent."),
            ("Recevoir un webhook FedaPay", "FedaPay", "Valider payload et état transaction.", "Commande/abonnement mis à jour."),
            ("Signer un upload média", "Utilisateur autorisé", "Demander une signature pour Cloudinary.", "Upload limité au contexte et au type autorisé."),
            ("Envoyer email transactionnel", "Système", "Déclencher template après événement métier.", "Email envoyé ou erreur journalisée sans casser la transaction."),
            ("Exécuter tâches cron", "Système", "Lancer rappels, expiration holds, payouts et récap.", "Traitement idempotent avec verrou cron."),
            ("Appliquer rate limits", "Système", "Compter les appels par scope/IP/utilisateur.", "429 avec retry-after lorsque la limite est dépassée."),
            ("Servir API mobile avec CORS", "Mobile", "Préflight puis requête avec origine autorisée.", "Headers CORS corrects et réponse métier."),
        ]),
    ]
    feature_rows = []
    counter = 1
    for domain, features in feature_groups:
        doc.add_heading(domain, level=3)
        for name, actor, description, result in features:
            feature_rows.append((f"F-{counter:03d}", name, f"{actor} — {description} Résultat : {result}"))
            counter += 1
    add_table(doc, ["ID", "Fonctionnalité", "Acteur, comportement et résultat"], feature_rows, [1100, 2200, 6060])

    doc.add_heading("6.2 Parcours utilisateur détaillés", level=2)
    journey_rows = [
        ("J-01", "Visiteur → client", "Découvrir puis acheter", "Accueil → recherche → détail événement → choix billet → checkout → paiement → confirmation → billet", "Compte ou checkout invité selon la règle; paiement accepté; ticket créé; email et notification."),
        ("J-02", "Client", "Suivre un événement", "Connexion → événement → intérêt → profil événements intéressés → rappel notification", "Interest persisté; retrait idempotent; rappel selon préférence."),
        ("J-03", "Client", "Suivre un organisateur", "Annuaire → profil organisateur → suivre → réglages d’alertes → liste suivis", "Follow et préférences alertes persistés; compteur cohérent."),
        ("J-04", "Client", "Acheter et utiliser un billet", "Commande → paiement → billet → QR/token → scanner à l’entrée", "Une seule validation; refus explicite si billet annulé, inconnu ou déjà utilisé."),
        ("J-05", "Client", "Transférer un billet", "Billets → assigner → destinataire reçoit invitation → accepter/refuser → propriétaire final", "Transfert atomique; impossibilité de double attribution."),
        ("J-06", "Client", "Revendre un billet", "Billet éligible → annonce → acheteur → checkout revente → transfert → clôture", "Prix/règles respectés; ticket change de propriétaire sans duplicat."),
        ("J-07", "Client", "Demander remboursement", "Commande → remboursement → motif → suivi → décision système/agent", "Une seule demande active; statut visible; notification à chaque étape."),
        ("J-08", "Client", "Créer une conversation", "Messages → nouveau → contact → conversation → message → réaction/lecture", "Conversation et message persistés; participant autorisé uniquement."),
        ("J-09", "Client", "Créer et gérer groupe", "Nouveau groupe → membres → nom/avatar → admin → mute/pin/leave", "Rôles groupe appliqués; départ protégé si dernier admin."),
        ("J-10", "Client", "Gérer son compte", "Profil → paramètres → identité/confidentialité/préférences → export ou suppression", "Données mises à jour; suppression contrôlée par dossier métier."),
        ("J-11", "Organisateur", "Devenir organisateur", "Inscription → candidature brouillon → documents → soumission → revue agent → activation", "Statut pending/approved/rejected; emails et accès studio synchronisés."),
        ("J-12", "Organisateur", "Créer et publier événement", "Studio → créer → billets/médias/staff → aperçu → publier → annuaire public", "Propriétaire et contraintes validés; événement découvrable."),
        ("J-13", "Organisateur", "Exploiter événement", "Événement → guestlist/staff/promo/playlist → ventes → stats → reversement", "Données ventes et droits agrégés; audit des actions."),
        ("J-14", "Organisateur", "Reporter/annuler", "Studio → événement → reporter/annuler → confirmation → impact clients", "Date validée ou annulation; remboursements/notifications idempotents."),
        ("J-15", "Prestataire", "Publier son offre", "Candidature → profil → catalogue → médias → abonnement/facturation → profil public", "Offre visible selon statut; média signé et catalogue cohérent."),
        ("J-16", "Agent", "Traiter une candidature", "Espace agent → dossiers → détail → note → approuver/refuser → notification", "Décision tracée; rôle/statut métier mis à jour."),
        ("J-17", "Agent", "Modérer un signalement", "Signalements → détail → vérifier contexte → traiter → note → notification", "Rapport handled; action adaptée et journalisée."),
        ("J-18", "Agent", "Traiter paiement/remboursement", "Paiements → alertes/refunds/payouts → vérifier → complete/settle/mark-paid", "État financier et audit mis à jour sans double traitement."),
        ("J-19", "Agent", "Gérer compte/suppression", "Comptes ou suppressions → détail → vérifier → disable/approve/reject", "Accès révoqué ou purge/anonymisation déclenchée."),
        ("J-20", "Système", "Webhook paiement", "Fournisseur → endpoint signé → validation → idempotence → commande/abonnement → email", "2xx si accepté; non-traité rejeté sans mutation partielle."),
    ]
    add_table(doc, ["ID", "Acteur", "Objectif", "Étapes", "Résultat / contrôles"], journey_rows, [700, 1500, 1500, 3300, 2360])

    doc.add_heading("7. API et intégrations", level=1)
    add_para(doc, f"Le code actuel expose {len(api)} fichiers de routes API. Les routes sont groupées ci-dessous par domaine. Les paramètres entre crochets sont dynamiques.")
    api_groups = [
        ("Auth, profil et compte", ["/api/auth", "/api/account", "/api/profil", "/api/profile", "/api/push"]),
        ("Découverte publique", ["/api/events", "/api/organizers", "/api/providers", "/api/search", "/api/preferences"]),
        ("Commandes et billets", ["/api/checkout", "/api/event-orders", "/api/orders", "/api/tickets", "/api/seat-holds", "/api/refund-link", "/api/resale-listings"]),
        ("Organisateur/prestataire", ["/api/organizer-events", "/api/organizers/me", "/api/providers/me", "/api/applications", "/api/subscriptions", "/api/my-staffed-events"]),
        ("Social et messagerie", ["/api/conversations", "/api/messages", "/api/friends", "/api/notifications", "/api/users"]),
        ("Playlist et avis", ["/api/events/[eventId]/playlist", "/api/reviews", "/api/providers/[providerId]/reviews"]),
        ("Agent", ["/api/agent", "/api/agent-sales"]),
        ("Paiements et webhooks", ["/api/stripe-webhook", "/api/webhooks/stripe", "/api/webhooks/fedapay", "/api/checkout/fedapay", "/api/seat-holds/fedapay"]),
        ("Tâches système", ["/api/cron", "/api/health", "/api/uploads"]),
    ]
    for title, prefixes in api_groups:
        doc.add_heading(title, level=2)
        matched = [e.replace("app/", "/").replace("/route.ts", "") for e in api if any(e.startswith("app" + p) for p in prefixes)]
        for entry in matched:
            add_bullet(doc, entry)
        if not matched:
            add_para(doc, "Aucune route détectée dans le périmètre courant.", color="6B7785")

    doc.add_heading("8. Règles métier et contrôles", level=1)
    for text in [
        "Une ressource privée est toujours filtrée par l’utilisateur courant et les participants autorisés.",
        "La propriété organisateur est vérifiée par organizerId/createdBy avant toute mutation d’événement.",
        "Les opérations de paiement, remboursement, reversement et annulation sont idempotentes et traçables.",
        "Les webhooks Stripe/FedaPay doivent vérifier leur signature et répondre en 2xx uniquement après traitement accepté.",
        "Un compte supprimé est purgé ou anonymisé selon son statut métier et les contraintes de dossier actif.",
        "Les tickets et tokens sont contrôlés contre la commande, l’événement, l’état du billet et les doublons de check-in.",
        "Les limites de pagination, recherche, upload et rate limit protègent les endpoints publics et connectés.",
        "Les modales et panneaux de détail ne doivent pas masquer le header, piéger le focus ou supprimer le retour clavier/mobile.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("9. Données métier", level=1)
    models = sorted(p.stem for p in (ROOT / "lib/models").glob("*.ts") if not p.name.startswith("__"))
    add_para(doc, "Collections principales détectées :")
    for i in range(0, len(models), 4):
        add_bullet(doc, ", ".join(models[i:i+4]))
    add_para(doc, "Les agrégats les plus importants sont User → OrganizerProfile/ProviderProfile, Event → Ticket/EventOrder, Conversation → Message, et les entités de modération/payment autour de Report, Review, Application, DeletionRequest, PaymentAlert et PayoutRequest.")

    doc.add_heading("10. Emails, notifications et événements externes", level=1)
    for text in [
        "Emails d’authentification : vérification, reset password, confirmation et sécurité du compte.",
        "Emails transactionnels : commande, billet, invitation, remboursement, revente, check-in et changements d’événement.",
        "Emails métier : candidature, modération, avis, suivi organisateur, payouts, abonnements et alertes agent.",
        "Notifications in-app : demandes d’amis, messages, réactions, événements suivis, modération et paiements.",
        "Images et médias : Cloudinary/signature upload, avec URLs publiques ou privées selon le document.",
        "Paiements : Stripe pour les flux carte/webhook et FedaPay pour les parcours régionaux configurés.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("11. Sécurité, confidentialité et accessibilité", level=1)
    for text in [
        "Auth.js gère les sessions et la protection CSRF des actions d’authentification.",
        "CORS autorise les origines web/mobile configurées et refuse les origines non prévues.",
        "Les réponses API distinguent auth_required, forbidden, not_found, invalid_input et rate_limited.",
        "Les données sensibles (documents, secrets, tokens, mots de passe) ne doivent jamais être exposées dans les payloads publics ou logs.",
        "Les écrans utilisent des composants accessibles : labels, états de chargement, erreurs visibles, focus, contraste et cibles tactiles.",
        "Les cookies publicitaires et le reciblage tiers sont exclus ; les choix analytics/préférences restent explicites.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("12. Recette fonctionnelle et critères d’acceptation", level=1)
    add_table(doc, ["Domaine", "Critère d’acceptation", "Preuve attendue"], [
        ("Public", "Toutes les pages publiques chargent, naviguent et présentent un état vide/erreur exploitable.", "Smoke routes + contrôle visuel responsive"),
        ("Auth", "Inscription, vérification, connexion, reset et déconnexion fonctionnent.", "Tests Auth.js et parcours QA"),
        ("Client", "Profil, intérêt, suivi, amis, messages, billets, commandes et notifications persistent.", "Parcours client authentifié"),
        ("Organisateur", "Création, gestion, publication, équipe, ventes, stats, playlist et payouts respectent la propriété.", "Tests intégration + compte organisateur"),
        ("Prestataire", "Profil, catalogue, médias, avis et abonnement sont cohérents.", "Tests intégration + compte prestataire"),
        ("Agent", "Files de modération, comptes, paiements, suppressions et homepage sont protégées et actionnables.", "Smoke agent + tests agent"),
        ("Paiement", "Checkout, webhooks, remboursements et idempotence ne génèrent pas de double traitement.", "Fixtures Stripe/FedaPay + logs"),
        ("Mobile", "Les appels mobiles trouvent une route/méthode et les exports Web/Android compilent.", "Contrat 207/207 + exports"),
        ("Production", "Health, pages publiques, API, CORS et webhook répondent aux statuts attendus.", "Smoke production + runtime API"),
    ], [1500, 5100, 2760])

    doc.add_heading("13. Commandes de vérification", level=1)
    for command in [
        "npm test -- --no-cache",
        "npm run lint:core",
        "npm run build",
        "npm run check:mobile-api",
        "npm run check:mobile-readiness",
        "npm run check:mobile-client-flow",
        "npm run check:mobile-agent-flow",
        "LIB_WEB_BASE_URL=https://liveinblack.com npm run check:mobile-api:runtime",
        "LIB_WEB_BASE_URL=https://liveinblack.com npm run check:mobile-web-cors",
        "LIB_WEB_BASE_URL=https://liveinblack.com npm run ops:smoke",
    ]:
        add_bullet(doc, command)

    doc.add_heading("14. Annexes — inventaire automatique du code", level=1)
    add_para(doc, f"Pages Next.js détectées : {len(pages)}. Routes API détectées : {len(api)}. Cet inventaire est généré depuis le code courant ; il doit être régénéré après l’ajout d’un domaine ou d’une route.")
    add_route_catalog(doc, "14.1 Pages détectées", pages)
    add_route_catalog(doc, "14.2 Routes API détectées", api)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
