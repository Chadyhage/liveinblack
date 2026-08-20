from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from build_functional_spec import (
    add_bullet,
    add_callout,
    add_number,
    add_para,
    add_table,
    style_run,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "LiveInBlack_Guide_Fonctionnel_Par_Role.docx"


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, before, after in [
        ("Heading 1", 18, "5B1A27", 16, 8),
        ("Heading 2", 14, "7E2438", 12, 6),
        ("Heading 3", 11.5, "3B1B23", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = 1.14

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run("LIVEINBLACK  |  Guide fonctionnel par rôle"), size=8.5, color="7C6A70")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer.add_run("Référence fonctionnelle • " + date.today().strftime("%d/%m/%Y")), size=8.5, color="7C6A70")


def feature(title, objective, actions, steps, result, cases=None):
    return {
        "title": title,
        "objective": objective,
        "actions": actions,
        "steps": steps,
        "result": result,
        "cases": cases or [],
    }


def add_feature(doc, code, item):
    doc.add_heading(f"{code} — {item['title']}", level=3)
    add_para(doc, "Objectif — " + item["objective"], bold_prefix="Objectif — ")
    add_para(doc, "Possibilités — " + item["actions"], bold_prefix="Possibilités — ")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run("Parcours pas à pas"), bold=True, color="7E2438")
    for number, step in enumerate(item["steps"], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.14
        style_run(p.add_run(f"{number}. "), bold=True, color="7E2438")
        style_run(p.add_run(step))
    add_callout(doc, "Résultat attendu", item["result"], fill="F8EEF1")
    if item["cases"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run("Cas particuliers et points d’attention"), bold=True, color="7E2438")
        for case in item["cases"]:
            add_bullet(doc, case)


COMMON = [
    feature(
        "Comprendre les rôles et choisir son espace",
        "Permettre à chaque personne d’identifier ce qu’elle peut accomplir dans LiveInBlack.",
        "Consulter la partie correspondant à son rôle principal et comprendre qu’un même compte peut évoluer vers plusieurs rôles.",
        ["Identifier son besoin principal : participer, organiser ou proposer un service.", "Lire la partie correspondant au rôle choisi.", "Suivre le parcours indiqué pour la fonctionnalité recherchée.", "Revenir aux parcours transversaux lorsqu’une action implique plusieurs rôles."],
        "La personne sait immédiatement où commencer et quelles étapes suivre.",
    ),
    feature(
        "Consulter les contenus publics",
        "Découvrir l’univers LiveInBlack avant de créer un compte.",
        "Voir les événements, organisateurs, prestataires, articles, pages d’aide et informations légales accessibles à tous.",
        ["Arriver sur l’accueil ou sur une page partagée.", "Utiliser le menu principal ou la recherche.", "Ouvrir une fiche pour consulter son contenu détaillé.", "Créer un compte uniquement lorsqu’une action personnelle le nécessite."],
        "Le visiteur obtient les informations essentielles sans obstacle inutile.",
        ["Les actions personnelles, comme acheter ou écrire à quelqu’un, demandent une connexion.", "Les menus publics restent visibles même lorsque la personne est déjà connectée."],
    ),
    feature(
        "Créer et sécuriser son accès",
        "Donner à chaque rôle un accès personnel et fiable.",
        "Créer un compte, confirmer son adresse, se connecter, se déconnecter et récupérer un accès oublié.",
        ["Choisir l’action de création de compte.", "Renseigner les informations demandées et accepter les règles applicables.", "Confirmer l’adresse de contact depuis le message reçu.", "Se connecter puis accéder à son espace personnel.", "En cas d’oubli, demander un nouveau mot de passe et en choisir un nouveau."],
        "Le compte est actif, identifiable et utilisable sur les espaces autorisés.",
        ["Un message clair explique si une confirmation est encore nécessaire.", "Une demande de récupération ne révèle jamais si une autre personne possède un compte."],
    ),
    feature(
        "Recevoir des confirmations utiles",
        "Rassurer la personne après une action importante.",
        "Recevoir une confirmation lisible après une inscription, un achat, une invitation, un changement important ou une demande d’assistance.",
        ["Réaliser l’action concernée.", "Voir immédiatement son nouvel état dans l’application.", "Recevoir, lorsque cela est pertinent, un message récapitulatif.", "Utiliser le bouton principal du message pour revenir au bon écran."],
        "La personne comprend ce qui a été enregistré et ce qu’elle doit faire ensuite.",
    ),
]


USER_SECTIONS = [
    ("A. Découvrir et choisir une sortie", [
        feature("Explorer l’accueil", "Obtenir une vue immédiate des sorties et tendances.", "Voir les événements mis en avant, les sélections, les catégories et les recommandations éditoriales.", ["Ouvrir l’accueil.", "Parcourir les sélections visibles sans changer de page.", "Choisir une carte selon la date, l’ambiance, le lieu ou l’intérêt.", "Ouvrir la fiche détaillée de l’événement."], "L’utilisateur passe rapidement de l’inspiration à une fiche exploitable."),
        feature("Rechercher rapidement", "Trouver un événement, un organisateur, un prestataire ou une personne à partir de quelques mots.", "Saisir un nom, un lieu, une ambiance ou un mot-clé et consulter des suggestions regroupées.", ["Ouvrir la recherche depuis l’en-tête.", "Saisir au moins quelques caractères.", "Lire les suggestions classées par type.", "Choisir un résultat ou afficher tous les résultats."], "Le bon contenu est accessible avec un minimum d’actions.", ["Un état vide propose de modifier les mots ou d’explorer les catégories."]),
        feature("Parcourir tous les événements", "Comparer l’offre disponible dans une liste claire.", "Voir les cartes d’événements, leur date, lieu, prix de départ, image, ambiance et disponibilité.", ["Ouvrir la liste des événements.", "Parcourir les résultats.", "Charger la suite ou changer de page.", "Ouvrir l’événement retenu."], "L’utilisateur dispose d’une vue d’ensemble cohérente et comparable."),
        feature("Filtrer et trier les événements", "Réduire la liste aux sorties réellement pertinentes.", "Filtrer par période, ville, catégorie, style musical, ambiance, prix ou disponibilité, puis trier les résultats.", ["Ouvrir les filtres.", "Choisir un ou plusieurs critères.", "Appliquer les critères et observer le nombre de résultats.", "Ajuster, retirer un filtre ou tout réinitialiser.", "Choisir l’ordre d’affichage souhaité."], "La sélection reflète les préférences exprimées et reste modifiable.", ["Les filtres actifs sont toujours visibles.", "Aucun résultat déclenche des suggestions de critères plus larges."]),
        feature("Consulter la fiche d’un événement", "Réunir toutes les informations nécessaires avant de décider.", "Voir titre, images, date, horaires, lieu, description, programme, artistes, ambiance, organisateur, catégories, billets, conditions et informations pratiques.", ["Ouvrir la fiche depuis une carte ou un lien partagé.", "Lire le résumé et les informations prioritaires.", "Parcourir les médias et les détails.", "Consulter les billets disponibles et les règles.", "Choisir une action : intérêt, partage, achat ou retour."], "L’utilisateur peut décider en connaissance de cause.", ["Un changement de date, un report, une annulation ou une rupture de stock est affiché clairement."]),
        feature("Partager un événement", "Permettre de faire découvrir une sortie à son entourage.", "Copier le lien ou utiliser les options de partage proposées.", ["Ouvrir la fiche de l’événement.", "Choisir Partager.", "Sélectionner une destination ou copier le lien.", "Envoyer le lien à la personne souhaitée."], "Le destinataire reçoit un accès direct à la fiche publique."),
        feature("Marquer un événement comme intéressant", "Garder une sortie en mémoire sans acheter immédiatement.", "Ajouter ou retirer un événement de la liste des événements intéressants.", ["Ouvrir une carte ou une fiche.", "Choisir l’action d’intérêt.", "Se connecter si nécessaire.", "Retrouver l’événement dans l’espace personnel.", "Retirer l’intérêt lorsqu’il n’est plus utile."], "L’événement est mémorisé et peut alimenter des rappels pertinents."),
        feature("Découvrir les organisateurs", "Identifier les structures derrière les événements.", "Parcourir l’annuaire, rechercher un nom et ouvrir un profil public.", ["Ouvrir la rubrique Organisateurs.", "Rechercher ou parcourir les profils.", "Comparer identité, ville, univers et événements.", "Ouvrir le profil retenu."], "L’utilisateur comprend qui organise et retrouve ses événements."),
        feature("Suivre un organisateur", "Ne pas manquer les prochaines annonces d’un organisateur apprécié.", "S’abonner ou se désabonner depuis le profil public et retrouver les organisateurs suivis.", ["Ouvrir le profil de l’organisateur.", "Choisir Suivre.", "Autoriser les alertes souhaitées.", "Consulter plus tard la liste des organisateurs suivis.", "Se désabonner à tout moment."], "Les nouveautés pertinentes peuvent être signalées à l’utilisateur."),
        feature("Découvrir les prestataires", "Trouver des professionnels adaptés à un besoin événementiel.", "Parcourir l’annuaire, rechercher, filtrer par catégorie ou zone et consulter les profils.", ["Ouvrir la rubrique Prestataires.", "Saisir le besoin ou choisir une catégorie.", "Comparer les profils, services, zones et avis.", "Ouvrir le profil le plus pertinent."], "L’utilisateur identifie un prestataire crédible et adapté."),
        feature("Lire le blog et les contenus d’aide", "S’informer et comprendre l’usage de la plateforme.", "Lire les articles, découvrir comment fonctionne le service et accéder aux réponses fréquentes.", ["Ouvrir Blog, À propos ou Aide.", "Choisir un sujet.", "Lire le contenu et ses illustrations.", "Utiliser les liens associés pour poursuivre un parcours."], "La personne obtient une information claire sans assistance individuelle."),
    ]),
    ("B. Gérer son compte et ses préférences", [
        feature("Créer un compte utilisateur", "Accéder aux actions personnelles de LiveInBlack.", "S’inscrire avec son identité, ses coordonnées et un mot de passe.", ["Choisir Créer un compte.", "Renseigner les champs obligatoires.", "Lire et accepter les conditions applicables.", "Valider l’inscription.", "Confirmer l’adresse de contact puis se connecter."], "Le compte utilisateur est créé et prêt à être personnalisé."),
        feature("Se connecter et se déconnecter", "Accéder à son espace privé puis le quitter en sécurité.", "Se connecter avec ses informations d’accès et fermer la session depuis le profil.", ["Choisir Connexion.", "Renseigner l’adresse et le mot de passe.", "Valider et retrouver le contenu demandé ou l’espace personnel.", "À la fin, ouvrir le menu du profil et choisir Déconnexion."], "L’accès personnel est ouvert puis fermé selon le choix de l’utilisateur."),
        feature("Récupérer un mot de passe oublié", "Retrouver l’accès sans assistance manuelle.", "Demander un lien temporaire et définir un nouveau mot de passe.", ["Choisir Mot de passe oublié.", "Saisir l’adresse utilisée pour le compte.", "Ouvrir le message reçu.", "Choisir et confirmer un nouveau mot de passe.", "Se reconnecter."], "L’ancien mot de passe est remplacé et l’accès est rétabli."),
        feature("Consulter et modifier son profil", "Maintenir une identité personnelle correcte.", "Voir et modifier nom, photo, biographie courte et informations affichées dans les espaces sociaux.", ["Ouvrir Mon profil.", "Choisir Modifier.", "Mettre à jour les informations souhaitées.", "Vérifier l’aperçu.", "Enregistrer."], "Le profil reflète les informations actuelles de l’utilisateur."),
        feature("Modifier ses coordonnées", "Garder des moyens de contact fiables.", "Changer adresse, téléphone ou nom avec confirmation lorsque nécessaire.", ["Ouvrir Paramètres.", "Choisir la coordonnée à modifier.", "Saisir la nouvelle valeur.", "Confirmer l’identité si demandé.", "Vérifier que la nouvelle valeur apparaît."], "Les coordonnées valides deviennent les références du compte."),
        feature("Modifier son mot de passe", "Renforcer ou renouveler la protection du compte.", "Remplacer le mot de passe depuis les paramètres.", ["Ouvrir Sécurité.", "Saisir le mot de passe actuel.", "Choisir un nouveau mot de passe conforme aux indications.", "Le confirmer.", "Enregistrer puis utiliser le nouveau mot de passe."], "Le nouveau mot de passe est actif."),
        feature("Renseigner ses préférences d’ambiance", "Personnaliser les découvertes et recommandations.", "Choisir styles musicaux, ambiances, catégories, villes ou autres centres d’intérêt.", ["Ouvrir Préférences.", "Sélectionner les thèmes pertinents.", "Retirer ceux qui ne correspondent plus.", "Enregistrer.", "Observer les sélections personnalisées proposées."], "Les suggestions deviennent plus proches des goûts déclarés."),
        feature("Gérer ses préférences de recherche", "Retrouver plus facilement les critères habituels.", "Mémoriser ou modifier les zones, périodes et types de sorties favoris.", ["Effectuer une recherche.", "Choisir de mémoriser les critères si l’option est proposée.", "Retrouver les préférences dans les paramètres.", "Modifier ou supprimer une préférence."], "Les recherches fréquentes demandent moins d’effort."),
        feature("Gérer confidentialité et visibilité", "Contrôler les informations visibles par les autres.", "Choisir la visibilité des éléments personnels et comprendre l’usage des données.", ["Ouvrir Confidentialité.", "Lire les catégories d’informations.", "Choisir le niveau de visibilité disponible.", "Enregistrer les changements.", "Vérifier le résultat sur le profil."], "Les choix de confidentialité sont explicites et respectés."),
        feature("Gérer les cookies", "Choisir les usages facultatifs tout en conservant le service essentiel.", "Accepter, refuser ou personnaliser les préférences d’ambiance et la mesure d’audience.", ["Lire le bandeau lors de la première visite.", "Accepter l’ensemble, refuser le facultatif ou personnaliser.", "Consulter le détail des catégories.", "Modifier plus tard le choix depuis la page Cookies."], "Le choix est mémorisé et peut être révisé."),
        feature("Exporter ses informations", "Obtenir une copie lisible de ses informations personnelles.", "Demander un export depuis les paramètres et récupérer le fichier préparé.", ["Ouvrir Confidentialité.", "Choisir Exporter mes informations.", "Confirmer la demande.", "Attendre la notification de disponibilité.", "Télécharger et conserver le document en sécurité."], "L’utilisateur reçoit une copie de ses informations.", ["Une nouvelle demande peut être temporairement limitée si un export est déjà en préparation."]),
        feature("Demander la suppression de son compte", "Permettre de quitter le service de façon maîtrisée.", "Lancer une demande, comprendre les conséquences et confirmer la décision.", ["Ouvrir Supprimer mon compte.", "Lire les conséquences sur billets, conversations et rôles.", "Régler les actions encore ouvertes si nécessaire.", "Confirmer la demande.", "Recevoir la confirmation et le délai annoncé."], "La demande est enregistrée et son état reste compréhensible.", ["Une confirmation renforcée évite une suppression accidentelle.", "Certaines informations peuvent être conservées uniquement lorsque des obligations l’exigent."]),
    ]),
    ("C. Réserver, payer et gérer ses billets", [
        feature("Choisir un billet", "Sélectionner l’offre adaptée à l’événement.", "Comparer les types de billets, prix, avantages, quantités et conditions.", ["Ouvrir la fiche de l’événement.", "Consulter les billets disponibles.", "Choisir un type et une quantité.", "Lire le récapitulatif avant de continuer."], "La sélection est claire avant toute validation."),
        feature("Réserver un billet gratuit", "Obtenir une place sans paiement.", "Choisir un billet gratuit et confirmer les informations du participant.", ["Sélectionner le billet gratuit.", "Se connecter si nécessaire.", "Vérifier l’identité et la quantité.", "Confirmer la réservation.", "Ouvrir le billet reçu."], "Le billet gratuit est immédiatement disponible dans Mes billets."),
        feature("Acheter un billet payant", "Finaliser une réservation avec un moyen de paiement proposé.", "Vérifier la commande, payer et recevoir la confirmation.", ["Choisir billets et quantités.", "Se connecter puis vérifier le récapitulatif, les frais et le total.", "Choisir le moyen de paiement.", "Suivre les instructions de validation.", "Revenir à la confirmation et vérifier Mes billets."], "La commande payée apparaît avec ses billets et son justificatif.", ["Un paiement interrompu laisse une information claire et permet de réessayer sans acheter deux fois."]),
        feature("Conserver sa sélection pendant le paiement", "Éviter qu’une place choisie disparaisse pendant une courte validation.", "Voir la durée restante et reprendre ou abandonner la commande.", ["Choisir une offre à disponibilité limitée.", "Lire le temps réservé annoncé.", "Finaliser le paiement avant l’expiration.", "Si le temps expire, revenir aux billets disponibles."], "La disponibilité et le temps restant sont compréhensibles."),
        feature("Comprendre la confirmation de paiement", "Savoir immédiatement si la commande a abouti.", "Voir le statut, le montant, l’événement, les billets et la prochaine action.", ["Terminer l’étape de paiement.", "Attendre la confirmation affichée.", "Vérifier le statut et le numéro de commande.", "Ouvrir les billets ou revenir à l’événement."], "Aucune ambiguïté ne subsiste sur l’issue du paiement."),
        feature("Consulter ses commandes", "Retrouver l’historique des achats.", "Voir les commandes par date, événement, montant et état.", ["Ouvrir l’espace personnel.", "Choisir Commandes ou Mes billets.", "Parcourir l’historique.", "Ouvrir une commande pour voir son détail."], "Chaque achat est traçable et compréhensible."),
        feature("Afficher un billet", "Présenter le droit d’entrée le jour de l’événement.", "Ouvrir le billet, son code de contrôle, le nom du participant et les informations pratiques.", ["Ouvrir Mes billets.", "Choisir l’événement.", "Ouvrir le billet concerné.", "Présenter le code au contrôle ou consulter les instructions."], "Le billet est rapidement accessible, y compris sur petit écran."),
        feature("Attribuer un billet à un proche", "Associer un billet acheté à la personne qui l’utilisera.", "Choisir un billet éligible, renseigner le destinataire et envoyer l’invitation.", ["Ouvrir la commande puis le billet.", "Choisir Attribuer ou Inviter.", "Rechercher le proche ou saisir ses coordonnées.", "Vérifier le destinataire.", "Envoyer l’invitation et suivre son état."], "Le billet attend l’acceptation du destinataire."),
        feature("Accepter ou refuser une invitation", "Permettre au destinataire de prendre une décision explicite.", "Consulter l’événement et accepter ou refuser le billet proposé.", ["Ouvrir la notification ou le lien reçu.", "Se connecter avec le bon compte.", "Lire l’événement, le billet et l’expéditeur.", "Choisir Accepter ou Refuser.", "Vérifier le résultat dans Mes billets."], "Le billet est attribué au destinataire uniquement après acceptation."),
        feature("Suivre ses invitations", "Connaître l’état des billets proposés à d’autres personnes.", "Voir les invitations en attente, acceptées, refusées ou annulées.", ["Ouvrir la commande concernée.", "Consulter les invitations envoyées.", "Lire l’état de chaque destinataire.", "Relancer personnellement ou annuler lorsque cela est permis."], "L’acheteur conserve une vision claire de chaque billet."),
        feature("Annuler une attribution", "Récupérer un billet encore éligible.", "Annuler une invitation en attente ou révoquer une attribution selon les règles annoncées.", ["Ouvrir le billet attribué.", "Vérifier que l’annulation est encore autorisée.", "Choisir Annuler l’attribution.", "Confirmer.", "Vérifier le retour du billet dans la commande."], "Le billet revient à son propriétaire initial lorsqu’il est encore récupérable."),
        feature("Quitter un billet attribué", "Permettre au destinataire de rendre un billet qu’il ne peut plus utiliser.", "Renoncer au billet avant son utilisation, selon les conditions applicables.", ["Ouvrir le billet reçu.", "Choisir Rendre ou Quitter le billet.", "Lire les conséquences.", "Confirmer.", "Vérifier qu’il n’apparaît plus comme utilisable."], "L’acheteur initial est informé et récupère la gestion du billet."),
        feature("Mettre un billet en revente", "Proposer un billet éligible à un autre utilisateur.", "Choisir le billet, consulter les conditions, définir le prix autorisé et publier l’annonce.", ["Ouvrir le billet inutilisé.", "Choisir Revendre.", "Lire l’éligibilité, le prix et la durée.", "Confirmer la mise en vente.", "Suivre l’état de l’annonce."], "Le billet est proposé sans pouvoir être utilisé simultanément.", ["Un billet déjà contrôlé, remboursé, attribué ou non éligible ne peut pas être revendu."]),
        feature("Acheter un billet en revente", "Acquérir une place remise en vente dans un cadre contrôlé.", "Voir l’offre de revente, le prix, les conditions et finaliser l’achat.", ["Ouvrir une offre de revente depuis l’événement.", "Vérifier le billet et le montant.", "Choisir Acheter.", "Finaliser le paiement.", "Retrouver le nouveau billet dans Mes billets."], "Le billet devient utilisable par le nouvel acheteur et ne l’est plus par l’ancien."),
        feature("Demander un remboursement", "Signaler une situation ouvrant droit à examen ou remboursement.", "Ouvrir la demande depuis la commande ou un lien reçu, choisir les billets et expliquer le motif.", ["Ouvrir la commande ou le lien de remboursement.", "Vérifier l’événement et les billets concernés.", "Choisir le motif et ajouter les informations utiles.", "Confirmer la demande.", "Suivre son état jusqu’à la décision."], "La demande est tracée et l’utilisateur connaît son statut.", ["Une demande n’équivaut pas toujours à une acceptation immédiate.", "En cas d’annulation ou de report, des consignes spécifiques peuvent s’appliquer."]),
    ]),
    ("D. Communauté, messagerie et participation", [
        feature("Rechercher une personne", "Retrouver un proche présent sur LiveInBlack.", "Rechercher par nom ou information autorisée et consulter un aperçu de profil.", ["Ouvrir la recherche de personnes.", "Saisir le nom recherché.", "Comparer les résultats sans exposer d’informations privées.", "Ouvrir le profil pertinent ou envoyer une demande d’ami."], "La bonne personne peut être identifiée dans le respect de sa visibilité."),
        feature("Envoyer une demande d’ami", "Créer une relation sociale consentie.", "Inviter une personne, suivre la demande et l’annuler avant réponse.", ["Ouvrir le profil de la personne.", "Choisir Ajouter.", "Vérifier l’envoi.", "Consulter les demandes envoyées.", "Annuler si nécessaire."], "La relation reste en attente tant que l’autre personne n’a pas accepté."),
        feature("Gérer les demandes reçues", "Décider qui rejoint son réseau.", "Accepter ou refuser chaque demande d’ami.", ["Ouvrir les demandes reçues.", "Consulter l’identité de l’expéditeur.", "Choisir Accepter ou Refuser.", "Vérifier la nouvelle relation ou la disparition de la demande."], "Seules les demandes acceptées deviennent des relations."),
        feature("Gérer sa liste d’amis", "Maintenir un réseau pertinent.", "Consulter ses amis, ouvrir un profil, démarrer une conversation ou retirer une relation.", ["Ouvrir Amis.", "Rechercher une personne dans la liste.", "Choisir l’action souhaitée.", "Confirmer toute suppression."], "La liste reflète les relations actuelles."),
        feature("Bloquer ou débloquer une personne", "Mettre fin aux interactions indésirables.", "Bloquer depuis un profil ou une conversation, consulter la liste des blocages et débloquer.", ["Ouvrir le profil ou les options de conversation.", "Choisir Bloquer.", "Lire les conséquences puis confirmer.", "Pour revenir sur la décision, ouvrir Personnes bloquées.", "Choisir Débloquer."], "Les interactions directes sont empêchées tant que le blocage reste actif."),
        feature("Signaler un comportement", "Alerter l’équipe LiveInBlack sur un contenu ou une personne problématique.", "Choisir un motif, ajouter un contexte et transmettre le signalement.", ["Ouvrir le menu du contenu, message ou profil.", "Choisir Signaler.", "Sélectionner le motif le plus précis.", "Ajouter une explication utile sans information inutile.", "Envoyer puis utiliser le blocage si besoin immédiat."], "Le signalement est enregistré pour examen.", ["Le signalement n’informe pas automatiquement la personne signalée de l’identité de l’auteur."]),
        feature("Démarrer une conversation privée", "Échanger directement avec une personne autorisée.", "Créer ou retrouver la conversation, envoyer un texte et suivre les réponses.", ["Ouvrir le profil ou la messagerie.", "Choisir Nouveau message.", "Sélectionner le destinataire.", "Écrire puis envoyer.", "Retrouver la conversation dans la boîte de réception."], "Une conversation unique et continue est disponible avec cette personne."),
        feature("Créer une conversation de groupe", "Coordonner plusieurs personnes autour d’une sortie.", "Choisir les membres, nommer le groupe et commencer l’échange.", ["Ouvrir Nouveau groupe.", "Sélectionner plusieurs participants.", "Donner un nom et éventuellement une image.", "Créer le groupe.", "Envoyer le premier message."], "Le groupe apparaît chez les membres ajoutés selon leurs autorisations."),
        feature("Administrer un groupe", "Maintenir une conversation collective organisée.", "Renommer, changer l’image, ajouter ou retirer des membres et attribuer les responsabilités disponibles.", ["Ouvrir les informations du groupe.", "Choisir l’élément à modifier.", "Effectuer le changement autorisé.", "Confirmer les actions sensibles.", "Vérifier le message d’information dans la conversation."], "La composition et l’identité du groupe restent à jour."),
        feature("Gérer une conversation", "Adapter la messagerie à ses besoins.", "Épingler, mettre en sourdine, masquer, vider l’historique local ou quitter un groupe.", ["Ouvrir les options de la conversation.", "Choisir l’action souhaitée.", "Définir une durée lorsque la mise en sourdine le propose.", "Confirmer une action irréversible.", "Vérifier le nouvel état dans la liste."], "La boîte de réception est organisée selon les préférences de l’utilisateur."),
        feature("Envoyer des contenus", "Partager plus qu’un simple texte.", "Envoyer image, vidéo, document ou autre contenu autorisé avec aperçu et progression.", ["Ouvrir une conversation.", "Choisir Joindre.", "Sélectionner le contenu.", "Vérifier l’aperçu et ajouter un texte si souhaité.", "Envoyer et attendre la confirmation."], "Le contenu apparaît dans la conversation avec un état compréhensible.", ["Un échec d’envoi conserve une possibilité de réessayer.", "Un contenu trop volumineux ou non accepté déclenche une explication claire."]),
        feature("Agir sur un message", "Corriger, organiser ou partager un message.", "Modifier son propre texte, supprimer selon les règles, transférer, mettre en favori ou réagir.", ["Appuyer sur le message concerné.", "Choisir l’action proposée.", "Modifier ou sélectionner la destination si nécessaire.", "Confirmer.", "Vérifier la marque de modification, suppression, favori ou réaction."], "L’action est visible et cohérente pour les personnes concernées."),
        feature("Créer un sondage", "Prendre une décision collective simplement.", "Poser une question, proposer des choix, définir les règles de vote et consulter les résultats.", ["Dans un groupe, choisir Créer un sondage.", "Saisir la question et les réponses possibles.", "Choisir si plusieurs réponses sont permises.", "Publier.", "Voter et consulter l’évolution des résultats."], "Le groupe dispose d’un résultat partagé et actualisé."),
        feature("Gérer ses notifications", "Rester informé sans être surchargé.", "Consulter les alertes, les marquer comme lues et choisir les catégories autorisées.", ["Ouvrir Notifications.", "Lire une alerte et accéder à son contenu.", "Marquer une alerte ou toutes les alertes comme lues.", "Ouvrir les préférences.", "Activer ou désactiver les catégories souhaitées."], "Les alertes importantes restent visibles selon les choix de l’utilisateur."),
        feature("Participer à la playlist d’un événement", "Contribuer à l’ambiance musicale lorsque l’organisateur l’autorise.", "Consulter la playlist, proposer un titre, aimer une proposition et voir le titre en cours.", ["Ouvrir la fiche ou le billet de l’événement.", "Accéder à Playlist.", "Rechercher un morceau puis le proposer.", "Aimer les propositions pertinentes.", "Retirer sa propre proposition si elle n’est plus souhaitée."], "La participation musicale est enregistrée dans le cadre fixé par l’organisateur."),
        feature("Publier un avis", "Partager une expérience utile après un événement ou une prestation.", "Attribuer une note, rédiger un commentaire, modifier ou supprimer son avis et signaler un avis inapproprié.", ["Ouvrir l’événement ou le prestataire éligible.", "Choisir Donner un avis.", "Saisir la note et un commentaire factuel.", "Publier.", "Modifier ou supprimer plus tard si nécessaire."], "L’avis apparaît selon les règles de publication et peut recevoir une réponse."),
        feature("Contacter l’assistance", "Obtenir de l’aide lorsque les contenus disponibles ne suffisent pas.", "Choisir un sujet, décrire le besoin et fournir les informations utiles.", ["Consulter d’abord l’aide.", "Ouvrir Contact.", "Choisir le motif.", "Décrire la situation avec les références nécessaires.", "Envoyer et conserver la confirmation."], "La demande est transmise avec suffisamment de contexte pour être traitée."),
    ]),
]


ORGANIZER_SECTIONS = [
    ("A. Devenir organisateur et construire son identité", [
        feature("Déposer une candidature organisateur", "Demander l’accès aux outils de création et de gestion d’événements.", "Présenter son activité, ses responsables, sa zone et les justificatifs demandés.", ["Ouvrir Devenir organisateur.", "Commencer la candidature.", "Renseigner identité, activité et coordonnées.", "Ajouter les justificatifs demandés.", "Relire puis envoyer pour examen."], "La candidature est enregistrée avec un état de suivi."),
        feature("Enregistrer une candidature en brouillon", "Compléter un dossier en plusieurs fois.", "Sauvegarder les informations déjà saisies puis reprendre plus tard.", ["Commencer la candidature.", "Remplir une première partie.", "Choisir Enregistrer et quitter.", "Revenir dans Ma candidature.", "Poursuivre jusqu’à l’envoi."], "Aucune information validée n’est perdue entre deux sessions."),
        feature("Suivre l’examen de sa candidature", "Comprendre où en est la demande et ce qui reste à faire.", "Voir si le dossier est en brouillon, envoyé, en cours d’examen, à compléter, accepté ou refusé.", ["Ouvrir Ma candidature.", "Lire l’état et la date de dernière évolution.", "Consulter les demandes de complément.", "Ajouter ou corriger les éléments demandés.", "Renvoyer le dossier si nécessaire."], "Le candidat connaît la prochaine action attendue."),
        feature("Créer son profil public organisateur", "Présenter une identité crédible aux participants.", "Renseigner nom public, description, ville, pays, zones, univers, photo et bannière.", ["Après acceptation, ouvrir le Studio organisateur.", "Compléter les informations principales.", "Ajouter photo et bannière.", "Choisir les catégories et zones.", "Prévisualiser puis publier le profil."], "Le profil public est complet, cohérent et trouvable."),
        feature("Modifier et prévisualiser son profil", "Maintenir une présentation actuelle.", "Corriger textes, coordonnées publiques, visuels et zones, puis vérifier le rendu public.", ["Ouvrir Profil organisateur.", "Choisir Modifier.", "Mettre à jour les éléments.", "Prévisualiser sur différents formats.", "Enregistrer et contrôler la page publique."], "Les visiteurs voient la version actualisée."),
        feature("Changer de rôle actif", "Passer de l’usage personnel au travail d’organisation sans créer un autre compte.", "Choisir le rôle utilisateur ou organisateur depuis l’espace personnel.", ["Ouvrir le sélecteur de rôle.", "Choisir Organisateur.", "Accéder au Studio et à ses menus.", "Revenir au rôle Utilisateur pour les achats personnels."], "La navigation et les actions s’adaptent au rôle choisi sans perdre les autres droits."),
    ]),
    ("B. Concevoir et publier un événement", [
        feature("Créer un événement en brouillon", "Commencer une nouvelle proposition sans la rendre publique.", "Créer une fiche de travail et l’enrichir progressivement.", ["Ouvrir Mes événements.", "Choisir Créer un événement.", "Donner un titre et les informations initiales.", "Enregistrer en brouillon.", "Reprendre depuis la liste des brouillons."], "L’événement existe dans l’espace organisateur mais n’est pas encore annoncé."),
        feature("Définir l’identité de l’événement", "Rendre l’offre immédiatement compréhensible.", "Renseigner titre, accroche, description, catégories, styles, ambiances, artistes et mots-clés.", ["Ouvrir le brouillon.", "Compléter le titre et le résumé.", "Décrire l’expérience et le programme.", "Choisir catégories, ambiances et styles.", "Ajouter artistes et informations utiles."], "La proposition possède un positionnement clair et recherchable."),
        feature("Définir date, horaires et lieu", "Donner des informations pratiques fiables.", "Choisir début, fin, fuseau horaire, adresse, ville, pays et précisions d’accès.", ["Ouvrir Informations pratiques.", "Choisir la date et les horaires.", "Renseigner le lieu et l’adresse.", "Ajouter les indications d’accès.", "Vérifier la cohérence puis enregistrer."], "La date et le lieu sont affichés sans ambiguïté."),
        feature("Ajouter les médias de l’événement", "Donner envie avec des visuels adaptés.", "Ajouter image principale, galerie et vidéos autorisées, les réordonner ou les retirer.", ["Ouvrir Médias.", "Choisir l’image principale.", "Ajouter les autres visuels.", "Recadrer ou réordonner si nécessaire.", "Prévisualiser le rendu public."], "La fiche dispose d’une identité visuelle cohérente sur tous les écrans."),
        feature("Créer les types de billets", "Définir les offres commerciales ou gratuites.", "Créer plusieurs billets avec nom, prix, quantité, période de vente, avantages et conditions.", ["Ouvrir Billetterie.", "Choisir Ajouter un type de billet.", "Renseigner nom, prix ou gratuité et quantité.", "Définir la période et les conditions.", "Enregistrer puis répéter pour les autres offres."], "Chaque offre est distincte et sa disponibilité est contrôlable."),
        feature("Vérifier l’événement avant publication", "Éviter les informations manquantes ou incohérentes.", "Consulter un résumé des éléments obligatoires et prévisualiser la page publique.", ["Ouvrir l’aperçu.", "Parcourir titre, date, lieu, médias, description et billets.", "Corriger chaque alerte affichée.", "Tester les principales actions en aperçu.", "Revenir au tableau de préparation."], "L’événement est prêt à être compris et acheté."),
        feature("Publier l’événement", "Rendre l’événement visible et ouvrir les actions autorisées.", "Confirmer la publication après les derniers contrôles.", ["Ouvrir le brouillon prêt.", "Choisir Publier.", "Lire le récapitulatif et les conséquences.", "Confirmer.", "Ouvrir la fiche publique pour contrôle final."], "L’événement devient visible selon sa configuration."),
        feature("Modifier un événement publié", "Maintenir les informations exactes après publication.", "Mettre à jour les champs encore modifiables et informer les personnes concernées lorsqu’un changement est important.", ["Ouvrir Mes événements puis l’événement.", "Choisir Modifier.", "Changer les éléments autorisés.", "Lire l’impact annoncé.", "Enregistrer et vérifier la fiche publique."], "La nouvelle information est visible et les participants sont informés lorsque nécessaire."),
        feature("Dépublier un événement", "Retirer temporairement une annonce lorsqu’aucune conséquence incompatible ne l’empêche.", "Masquer la fiche de la découverte publique tout en conservant le brouillon de travail.", ["Ouvrir l’événement.", "Choisir Dépublier.", "Lire les effets sur les ventes et les liens existants.", "Confirmer.", "Corriger puis republier lorsque l’événement est prêt."], "L’événement n’est plus proposé au public mais reste gérable."),
        feature("Reporter un événement", "Changer la date en protégeant l’information des participants.", "Définir la nouvelle date, expliquer le report et transmettre les options disponibles.", ["Ouvrir Gestion de l’événement.", "Choisir Reporter.", "Renseigner la nouvelle date et le message.", "Vérifier les conséquences sur billets et remboursements.", "Confirmer et suivre les réactions des participants."], "La fiche affiche le report et les détenteurs de billets reçoivent une information claire."),
        feature("Annuler un événement", "Clôturer un événement qui ne peut plus avoir lieu.", "Déclarer l’annulation, expliquer la situation et organiser les suites pour les commandes.", ["Ouvrir Gestion de l’événement.", "Choisir Annuler.", "Lire toutes les conséquences.", "Saisir le message destiné aux participants.", "Confirmer puis suivre les remboursements et demandes."], "L’événement est marqué annulé et aucune nouvelle vente normale n’est possible."),
    ]),
    ("C. Promouvoir et animer la communauté", [
        feature("Suivre l’audience de son profil", "Comprendre l’intérêt suscité par l’organisateur.", "Voir le nombre d’abonnés et l’évolution des interactions disponibles.", ["Ouvrir le Studio.", "Consulter l’indicateur d’abonnés.", "Comparer les périodes.", "Identifier les annonces ayant créé de l’intérêt."], "L’organisateur dispose de repères simples pour orienter sa communication."),
        feature("Informer ses abonnés", "Faire connaître un nouvel événement ou un changement important.", "Déclencher les informations prévues lors d’une publication, d’un report ou d’une annulation.", ["Réaliser l’action importante sur l’événement.", "Vérifier le message proposé.", "Confirmer l’information des personnes concernées.", "Suivre les retours ou demandes."], "Les abonnés et participants reçoivent une information adaptée à leur relation avec l’événement."),
        feature("Mettre un événement en avant", "Accroître temporairement la visibilité d’un événement éligible.", "Vérifier les offres disponibles, choisir une mise en avant, payer et suivre sa période active.", ["Ouvrir Promotion depuis l’événement.", "Vérifier l’éligibilité et les emplacements proposés.", "Choisir la durée ou l’offre.", "Vérifier le montant puis payer.", "Suivre l’état et la date de fin."], "La mise en avant est identifiable et active pendant la période annoncée."),
        feature("Activer la playlist participative", "Impliquer les participants dans l’ambiance musicale.", "Ouvrir les suggestions, fixer les règles et modérer les morceaux.", ["Ouvrir Playlist dans l’événement.", "Activer la participation.", "Définir les règles visibles.", "Consulter les propositions et leur popularité.", "Retirer les propositions inadaptées si nécessaire."], "Les participants peuvent proposer des titres dans un cadre maîtrisé."),
        feature("Afficher le titre en cours", "Donner une information en direct aux participants.", "Choisir un morceau de la playlist comme titre actuellement joué.", ["Ouvrir la playlist de gestion.", "Repérer le morceau concerné.", "Choisir En cours de lecture.", "Changer ou retirer l’indication au moment voulu."], "Les participants voient le titre en cours lorsque cette information est activée."),
    ]),
    ("D. Vendre, accueillir et piloter l’événement", [
        feature("Suivre les ventes", "Connaître la performance commerciale de chaque événement.", "Voir commandes, billets vendus, billets restants, chiffre d’affaires et évolution dans le temps.", ["Ouvrir l’événement dans le Studio.", "Choisir Ventes ou Statistiques.", "Sélectionner la période.", "Comparer les types de billets.", "Utiliser les résultats pour ajuster la communication."], "L’organisateur comprend l’état des ventes sans calcul manuel."),
        feature("Consulter les commandes et participants", "Retrouver une commande ou un détenteur de billet.", "Rechercher, filtrer et ouvrir le détail des achats liés à l’événement.", ["Ouvrir Commandes.", "Rechercher par nom, référence ou état.", "Ouvrir la ligne concernée.", "Consulter billets, paiement, attribution et contrôle."], "L’équipe retrouve rapidement l’information nécessaire à l’assistance."),
        feature("Préparer la liste d’accueil", "Disposer d’une vision opérationnelle des personnes attendues.", "Consulter les billets valides, leurs détenteurs et leur état de contrôle.", ["Ouvrir Participants.", "Filtrer par type de billet ou état.", "Vérifier les noms et quantités.", "Utiliser la liste pendant l’accueil selon les droits accordés."], "L’équipe sait qui est attendu et qui est déjà entré."),
        feature("Constituer l’équipe de l’événement", "Déléguer les tâches sans donner plus de droits que nécessaire.", "Ajouter des membres, choisir leur mission et gérer leur accès aux événements concernés.", ["Ouvrir Équipe.", "Choisir Ajouter un membre.", "Identifier la personne.", "Attribuer la mission : vente, accueil, contrôle ou autre rôle proposé.", "Définir l’événement et la période puis confirmer."], "Chaque membre voit uniquement les outils nécessaires à sa mission."),
        feature("Planifier les créneaux de travail", "Organiser la présence des vendeurs et contrôleurs.", "Créer, attribuer et consulter les créneaux liés à un événement.", ["Ouvrir Équipe ou Créneaux.", "Choisir l’événement et la mission.", "Définir date, heure et personne.", "Confirmer.", "Vérifier que le membre retrouve le créneau dans Mes missions."], "Les responsabilités et horaires sont clairs avant l’ouverture."),
        feature("Vendre sur place", "Enregistrer des ventes réalisées à l’entrée ou à un point physique.", "Choisir des billets, la quantité, le mode de règlement et remettre la confirmation au participant.", ["Ouvrir Vente sur place pour l’événement.", "Choisir le billet et la quantité.", "Renseigner les informations nécessaires du participant.", "Choisir le règlement réellement reçu.", "Confirmer et remettre le billet ou justificatif."], "La vente est intégrée au suivi général et le billet peut être contrôlé."),
        feature("Gérer une commande sur place", "Corriger la composition d’une commande avant sa clôture.", "Ajouter, retirer ou modifier les quantités tant que l’état le permet.", ["Ouvrir la commande en cours.", "Choisir la ligne à modifier.", "Ajuster quantité ou offre.", "Vérifier le nouveau total.", "Confirmer la commande finale."], "La commande correspond à la vente réellement effectuée."),
        feature("Contrôler les billets à l’entrée", "Valider rapidement les droits d’accès.", "Lire le code d’un billet, voir son état et confirmer l’entrée.", ["Ouvrir Contrôle pour le bon événement.", "Autoriser l’accès nécessaire à l’appareil si demandé.", "Présenter ou saisir le code du billet.", "Lire le résultat : valide, déjà utilisé, non reconnu ou non valable.", "Confirmer l’entrée uniquement si le résultat l’autorise."], "Chaque entrée validée est comptabilisée et un billet ne peut pas servir deux fois.", ["Une procédure de recherche manuelle aide lorsque le code ne peut pas être lu."]),
        feature("Suivre la fréquentation en direct", "Connaître le niveau d’arrivée du public.", "Voir le nombre de billets contrôlés et le comparer aux ventes ou à la capacité.", ["Ouvrir Statistiques de l’événement.", "Consulter la fréquentation.", "Comparer entrées, billets vendus et capacité.", "Actualiser la répartition par type de billet."], "L’équipe dispose d’un aperçu utile pour adapter l’accueil."),
        feature("Consulter le bilan de l’événement", "Évaluer les résultats après la clôture.", "Voir ventes, entrées, répartition des billets, revenus, remboursements et indicateurs d’intérêt.", ["Ouvrir un événement terminé.", "Choisir Bilan ou Statistiques.", "Sélectionner la période complète.", "Lire les indicateurs et comparaisons.", "Conserver les conclusions pour les prochains événements."], "L’organisateur dispose d’un résumé exploitable de la performance."),
    ]),
    ("E. Gérer les revenus et versements", [
        feature("Configurer son profil de versement", "Indiquer où recevoir les sommes dues.", "Renseigner les informations du bénéficiaire et le moyen de versement proposé.", ["Ouvrir Finances.", "Choisir Configurer les versements.", "Renseigner les informations demandées.", "Vérifier attentivement le bénéficiaire.", "Confirmer et attendre la validation annoncée."], "Le profil de versement est prêt ou affiche précisément les éléments à compléter."),
        feature("Consulter son solde", "Comprendre les montants disponibles, en attente ou déjà versés.", "Voir la répartition des revenus et les retenues clairement expliquées.", ["Ouvrir Finances.", "Consulter le résumé du solde.", "Choisir un événement ou une période.", "Ouvrir le détail des mouvements."], "Chaque montant est rattaché à une origine et à un état."),
        feature("Demander un versement", "Recevoir un montant devenu disponible.", "Choisir le montant autorisé, vérifier le destinataire et confirmer la demande.", ["Ouvrir Solde disponible.", "Choisir Demander un versement.", "Vérifier montant, bénéficiaire et éventuels frais.", "Confirmer.", "Suivre l’état jusqu’à la réception."], "La demande est enregistrée avec un état et une date estimée."),
        feature("Suivre l’historique financier", "Retrouver toutes les opérations importantes.", "Parcourir versements, ventes, remboursements, ajustements et mises en avant payées.", ["Ouvrir Historique.", "Filtrer par période, événement ou type.", "Ouvrir une opération.", "Lire montant, date, état et référence fonctionnelle."], "L’organisateur peut expliquer l’évolution de son solde."),
    ]),
]


PROVIDER_SECTIONS = [
    ("A. Devenir prestataire et publier son profil", [
        feature("Déposer une candidature prestataire", "Demander l’accès à l’espace professionnel de services.", "Présenter son activité, ses catégories, ses zones d’intervention et ses justificatifs.", ["Ouvrir Devenir prestataire.", "Commencer la candidature.", "Renseigner identité, activité et coordonnées.", "Choisir catégories et zones.", "Ajouter les justificatifs puis envoyer."], "La candidature est enregistrée et visible dans le suivi."),
        feature("Compléter la candidature en plusieurs fois", "Éviter de perdre un dossier incomplet.", "Enregistrer un brouillon et reprendre les étapes restantes.", ["Commencer le dossier.", "Compléter les parties disponibles.", "Enregistrer le brouillon.", "Revenir dans Ma candidature.", "Relire puis envoyer lorsque tout est prêt."], "Les informations déjà saisies restent disponibles."),
        feature("Suivre la décision", "Savoir si l’activité peut être publiée et quoi corriger.", "Voir l’état, les observations et les demandes de complément.", ["Ouvrir Ma candidature.", "Lire l’état actuel.", "Consulter toute demande de précision.", "Ajouter les éléments manquants.", "Renvoyer le dossier pour examen."], "Le candidat connaît la décision ou la prochaine étape."),
        feature("Créer son profil public prestataire", "Présenter clairement l’entreprise et son expertise.", "Renseigner nom, description, catégories, ville, pays, zones, photo, bannière et moyens de contact publics.", ["Après acceptation, ouvrir l’espace Prestataire.", "Compléter l’identité et la description.", "Choisir catégories et zones.", "Ajouter les visuels.", "Prévisualiser puis publier."], "Le profil apparaît dans l’annuaire avec une présentation complète."),
        feature("Modifier et prévisualiser son profil", "Maintenir les informations professionnelles exactes.", "Changer textes, coordonnées publiques, zones et médias.", ["Ouvrir Profil prestataire.", "Choisir Modifier.", "Mettre à jour les éléments.", "Prévisualiser.", "Enregistrer et contrôler la page publique."], "La version publique reflète l’offre actuelle."),
        feature("Changer de rôle actif", "Passer de l’usage personnel à l’activité de prestataire.", "Choisir le rôle Prestataire ou Utilisateur depuis le même compte.", ["Ouvrir le sélecteur de rôle.", "Choisir Prestataire.", "Accéder au profil et au catalogue.", "Revenir au rôle Utilisateur pour les sorties personnelles."], "Les menus s’adaptent au rôle sans créer de compte séparé."),
    ]),
    ("B. Construire et maintenir son catalogue", [
        feature("Créer une prestation", "Présenter une offre précise aux visiteurs.", "Définir nom, catégorie, description, prix indicatif, unité, zone et conditions.", ["Ouvrir Mes services.", "Choisir Ajouter un service.", "Renseigner les informations principales.", "Ajouter prix ou mode d’estimation et conditions.", "Enregistrer puis prévisualiser."], "La prestation est compréhensible et rattachée au bon profil."),
        feature("Illustrer une prestation", "Montrer la qualité et le style du service.", "Ajouter une image principale et une galerie de réalisations.", ["Ouvrir le service.", "Choisir Médias.", "Ajouter des visuels représentatifs et autorisés.", "Réordonner et choisir l’image principale.", "Vérifier l’aperçu."], "La fiche de service dispose d’éléments visuels crédibles."),
        feature("Définir prix et conditions", "Donner un cadre de décision sans ambiguïté.", "Afficher un prix fixe, un prix de départ ou une indication sur devis, ainsi que les inclusions et limites.", ["Ouvrir Tarification.", "Choisir la forme de prix.", "Renseigner le montant et l’unité lorsqu’ils s’appliquent.", "Décrire ce qui est inclus et les éventuels frais.", "Enregistrer."], "Le visiteur comprend le niveau de prix et ce qu’il couvre."),
        feature("Définir zones et disponibilité", "Éviter les demandes incompatibles avec la capacité du prestataire.", "Préciser les villes ou zones couvertes et l’état général de disponibilité.", ["Ouvrir le service ou le profil.", "Choisir les zones couvertes.", "Indiquer les limites de déplacement.", "Mettre à jour la disponibilité affichée.", "Enregistrer."], "Les demandes potentielles correspondent mieux au périmètre réel."),
        feature("Publier ou masquer une prestation", "Contrôler ce que les visiteurs peuvent découvrir.", "Rendre un service visible, le masquer temporairement ou le republier.", ["Ouvrir Mes services.", "Choisir la prestation.", "Vérifier son niveau de complétude.", "Choisir Publier ou Masquer.", "Contrôler le résultat sur le profil public."], "Le catalogue public contient uniquement les offres actuellement proposées."),
        feature("Modifier une prestation", "Maintenir l’offre à jour.", "Corriger description, prix, conditions, zones ou médias.", ["Ouvrir le service.", "Choisir Modifier.", "Changer les éléments nécessaires.", "Prévisualiser.", "Enregistrer."], "Les visiteurs consultent les informations actualisées."),
        feature("Archiver ou supprimer une prestation", "Retirer une ancienne offre tout en évitant une perte accidentelle.", "Archiver une prestation pour la conserver hors publication ou la supprimer lorsqu’elle n’a plus d’utilité.", ["Ouvrir la prestation.", "Choisir Archiver ou Supprimer.", "Lire la différence et les conséquences.", "Confirmer.", "Vérifier la liste des services actifs ou archivés."], "Le catalogue actif reste propre et la décision est maîtrisée."),
    ]),
    ("C. Être découvert, contacté et évalué", [
        feature("Apparaître dans l’annuaire", "Être trouvé par les personnes recherchant un service.", "Présenter un profil publié et suffisamment complet dans les catégories et zones choisies.", ["Compléter le profil.", "Publier au moins une prestation.", "Vérifier catégories et zones.", "Contrôler l’apparition dans l’annuaire."], "Le prestataire apparaît dans les recherches pertinentes."),
        feature("Recevoir une prise de contact", "Transformer une consultation de profil en échange professionnel.", "Recevoir les coordonnées ou le message transmis par une personne intéressée, selon les options proposées.", ["Le visiteur ouvre le profil.", "Il choisit Contacter.", "Il décrit son besoin et fournit les informations nécessaires.", "Le prestataire reçoit la demande.", "Le prestataire répond par le canal indiqué."], "La demande contient un contexte suffisant pour engager l’échange."),
        feature("Consulter les avis reçus", "Comprendre la perception des clients.", "Voir notes, commentaires, dates et prestations concernées.", ["Ouvrir Avis.", "Consulter la note globale.", "Parcourir les avis récents.", "Filtrer ou ouvrir le détail si proposé."], "Le prestataire identifie ses points forts et axes d’amélioration."),
        feature("Répondre à un avis", "Apporter un complément public et professionnel.", "Publier une réponse courtoise rattachée à l’avis.", ["Ouvrir l’avis.", "Choisir Répondre.", "Rédiger une réponse factuelle sans donnée privée.", "Relire.", "Publier."], "La réponse apparaît avec l’avis et améliore la compréhension du contexte."),
        feature("Signaler un avis inapproprié", "Demander l’examen d’un contenu contraire aux règles.", "Choisir un motif précis et expliquer le problème.", ["Ouvrir l’avis.", "Choisir Signaler.", "Sélectionner le motif.", "Ajouter un contexte factuel.", "Envoyer et suivre l’information disponible."], "La demande d’examen est enregistrée sans supprimer automatiquement une critique légitime."),
    ]),
    ("D. Gérer l’abonnement et la continuité de visibilité", [
        feature("Choisir une formule", "Sélectionner le niveau d’accès adapté à l’activité.", "Comparer contenu, prix, durée et conditions de chaque formule proposée.", ["Ouvrir Abonnement.", "Comparer les formules.", "Choisir la zone de facturation si demandé.", "Lire le récapitulatif.", "Continuer vers le règlement."], "Le prestataire sait exactement ce qu’il choisit et pour combien de temps."),
        feature("Souscrire et payer", "Activer ou prolonger l’accès professionnel.", "Valider la formule, régler et recevoir une confirmation.", ["Sélectionner la formule.", "Vérifier le montant et la période.", "Choisir le moyen de paiement.", "Valider le règlement.", "Revenir à la confirmation et vérifier l’état de l’abonnement."], "L’abonnement actif affiche sa période et ses avantages."),
        feature("Consulter l’état de l’abonnement", "Connaître la formule, la période et la prochaine échéance.", "Voir état actif, en attente, expiré ou annulé ainsi que les dates importantes.", ["Ouvrir Abonnement.", "Lire la formule actuelle.", "Vérifier date de début, fin et renouvellement.", "Consulter les actions disponibles."], "Le prestataire sait si sa visibilité et ses outils sont maintenus."),
        feature("Renouveler ou changer de formule", "Adapter la continuité du service à l’évolution de l’activité.", "Prolonger une formule ou choisir une autre offre selon les conditions affichées.", ["Ouvrir la formule actuelle.", "Choisir Renouveler ou Changer.", "Comparer l’impact et la nouvelle période.", "Confirmer le choix et le règlement éventuel.", "Vérifier le nouvel état."], "La nouvelle période ou formule est clairement enregistrée."),
        feature("Annuler le renouvellement", "Éviter une nouvelle période non souhaitée.", "Désactiver le prochain renouvellement tout en conservant l’accès jusqu’à la date annoncée.", ["Ouvrir Abonnement.", "Choisir Annuler le renouvellement.", "Lire la date de fin et les conséquences.", "Confirmer.", "Vérifier la mention de fin programmée."], "Aucun nouveau renouvellement n’est attendu après la période déjà acquise."),
        feature("Réagir à une expiration ou un paiement non abouti", "Rétablir l’accès professionnel sans confusion.", "Comprendre l’état, corriger le moyen de paiement ou choisir une nouvelle formule.", ["Ouvrir l’alerte ou la rubrique Abonnement.", "Lire la cause et les conséquences sur la visibilité.", "Mettre à jour les informations nécessaires.", "Relancer le règlement ou choisir une formule.", "Vérifier le retour à l’état actif."], "Le prestataire sait comment restaurer sa visibilité."),
    ]),
]


USER_JOURNEYS = [
    ("Découvrir → acheter → participer", "Accueil → recherche ou sélection → fiche événement → choix du billet → règlement → confirmation → Mes billets → présentation du billet à l’entrée."),
    ("Suivre → être informé → réserver", "Profil organisateur → Suivre → choix des alertes → réception d’une nouveauté → fiche événement → réservation."),
    ("Acheter pour un proche", "Commande → billet → attribution → invitation → acceptation par le proche → billet dans son espace."),
    ("Revendre un billet", "Mes billets → billet éligible → conditions de revente → publication → achat par une autre personne → transfert définitif."),
    ("Organiser une sortie entre amis", "Recherche de personnes → relations → création d’un groupe → partage d’événement → sondage → réservation individuelle."),
    ("Demander de l’aide", "Aide → recherche de réponse → Contact → description du besoin → confirmation → suivi de la réponse."),
]

ORGANIZER_JOURNEYS = [
    ("Devenir organisateur", "Candidature → brouillon → justificatifs → envoi → examen → compléments éventuels → acceptation → création du profil public."),
    ("Lancer un événement", "Brouillon → identité → date et lieu → médias → billetterie → aperçu → publication → promotion → suivi des ventes."),
    ("Préparer le jour J", "Participants → équipe → créneaux → vente sur place → contrôle des billets → suivi des entrées."),
    ("Gérer un changement majeur", "Événement publié → report ou annulation → message explicatif → information des participants → traitement des demandes."),
    ("Clôturer et encaisser", "Fin de l’événement → bilan → vérification des mouvements → solde disponible → demande de versement → suivi."),
]

PROVIDER_JOURNEYS = [
    ("Devenir prestataire", "Candidature → brouillon → justificatifs → envoi → examen → acceptation → profil public."),
    ("Publier son offre", "Profil → catégories et zones → création d’une prestation → prix et conditions → médias → publication → contrôle dans l’annuaire."),
    ("Transformer une visite en contact", "Recherche publique → profil → service → avis → prise de contact → réception de la demande → réponse du prestataire."),
    ("Maintenir sa visibilité", "État de l’abonnement → échéance → renouvellement ou changement → paiement → confirmation → profil maintenu actif."),
]


def add_role_overview(doc, title, promise, sections, journey_rows):
    doc.add_heading(title, level=1)
    add_callout(doc, "Finalité du rôle", promise, fill="F8EEF1")
    rows = []
    total = 0
    for section_title, items in sections:
        rows.append((section_title, len(items), ", ".join(item["title"] for item in items[:3]) + ("…" if len(items) > 3 else "")))
        total += len(items)
    add_table(doc, ["Domaine", "Nombre", "Exemples"], rows, [3200, 1250, 4910])
    add_para(doc, f"Cette partie décrit {total} fonctionnalités principales. Chacune contient un objectif, les possibilités offertes, le parcours détaillé, le résultat attendu et les cas particuliers utiles.")
    doc.add_heading("Parcours de référence", level=2)
    add_table(doc, ["Parcours", "Enchaînement fonctionnel"], journey_rows, [2850, 6510])


def add_sections(doc, prefix, sections):
    index = 1
    for section_title, items in sections:
        doc.add_heading(section_title, level=2)
        for item in items:
            add_feature(doc, f"{prefix}-{index:03d}", item)
            index += 1


def add_page_break(doc):
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run("LIVEINBLACK"), size=14, bold=True, color="C83756")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    style_run(p.add_run("Guide fonctionnel complet"), size=30, bold=True, color="3B1B23")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(17)
    style_run(p.add_run("Toutes les fonctionnalités et tous les parcours, organisés par rôle"), size=15, color="7E5E66")
    add_table(doc, ["Partie 1", "Partie 2", "Partie 3", "Lecture"], [["Utilisateur", "Organisateur", "Prestataire", "100 % fonctionnelle et non technique"]], [2200, 2200, 2200, 2760])
    add_callout(doc, "Objet du document", "Expliquer, dans un langage accessible à tous, ce que chaque rôle peut faire dans LiveInBlack et le chemin exact pour y parvenir. Le document décrit l’expérience attendue, sans entrer dans la fabrication du produit.", fill="F8EEF1")
    add_para(doc, "Version — 1.0", bold_prefix="Version — ")
    add_para(doc, "Date — " + date.today().strftime("%d/%m/%Y"), bold_prefix="Date — ")
    add_para(doc, "Public — Direction, produit, opérations, support, partenaires, testeurs et futurs utilisateurs.", bold_prefix="Public — ")

    add_page_break(doc)
    doc.add_heading("Sommaire", level=1)
    add_table(doc, ["Partie", "Contenu"], [
        ("0. Mode d’emploi", "Principes communs, niveaux de lecture et règles d’expérience."),
        ("1. Utilisateur", "Découverte, compte, billets, paiements, invitations, revente, communauté, messagerie et participation."),
        ("2. Organisateur", "Candidature, profil, création d’événements, promotion, ventes, accueil, statistiques et versements."),
        ("3. Prestataire", "Candidature, profil, catalogue, visibilité, contacts, avis et abonnement."),
        ("4. Parcours transversaux", "Passage d’un rôle à l’autre, confirmations, protection, assistance et qualité d’usage."),
        ("5. Index fonctionnel", "Repérage rapide des fonctionnalités par rôle."),
    ], [2600, 6760])

    doc.add_heading("0. Mode d’emploi", level=1)
    add_para(doc, "Une fonctionnalité correspond à une capacité concrète offerte à une personne. Le parcours indique la suite d’actions permettant d’obtenir le résultat attendu. Les cas particuliers précisent les situations qui demandent une information, une confirmation ou une restriction supplémentaire.")
    add_table(doc, ["Repère", "Signification"], [
        ("Objectif", "Le besoin auquel la fonctionnalité répond."),
        ("Possibilités", "Ce que la personne peut réellement accomplir."),
        ("Parcours pas à pas", "La suite ordonnée d’actions à réaliser."),
        ("Résultat attendu", "Ce qui doit être vrai à la fin du parcours."),
        ("Cas particuliers", "Les exceptions, restrictions ou messages indispensables."),
    ], [2200, 7160])

    doc.add_heading("0.1 Principes communs", level=2)
    for i, item in enumerate(COMMON, 1):
        add_feature(doc, f"C-{i:03d}", item)

    add_page_break(doc)
    add_role_overview(doc, "1. Partie utilisateur", "Découvrir des expériences, réserver, gérer ses billets, participer à la communauté et garder le contrôle de son compte.", USER_SECTIONS, USER_JOURNEYS)
    add_sections(doc, "U", USER_SECTIONS)

    add_page_break(doc)
    add_role_overview(doc, "2. Partie organisateur", "Concevoir, vendre, promouvoir et exploiter des événements, depuis la candidature jusqu’au bilan et au versement des revenus.", ORGANIZER_SECTIONS, ORGANIZER_JOURNEYS)
    add_sections(doc, "O", ORGANIZER_SECTIONS)

    add_page_break(doc)
    add_role_overview(doc, "3. Partie prestataire", "Présenter une activité de service, publier un catalogue crédible, être découvert, recevoir des contacts et maintenir sa visibilité professionnelle.", PROVIDER_SECTIONS, PROVIDER_JOURNEYS)
    add_sections(doc, "P", PROVIDER_SECTIONS)

    add_page_break(doc)
    doc.add_heading("4. Parcours transversaux et règles d’expérience", level=1)
    transversal = [
        ("Un compte, plusieurs rôles", "Une personne conserve un seul accès et choisit le rôle actif. Le changement de rôle adapte les menus et les actions sans effacer les informations des autres rôles."),
        ("Connexion au bon moment", "Les contenus publics restent consultables. La connexion est demandée lorsqu’une action devient personnelle, financière, sociale ou professionnelle."),
        ("Droits liés au rôle", "Une action de gestion n’est proposée qu’à la personne concernée : propriétaire d’un billet, responsable d’un événement, membre autorisé d’une équipe ou propriétaire d’un profil."),
        ("Confirmation des décisions sensibles", "Paiement, annulation, suppression, blocage, transfert, revente et changement majeur demandent un récapitulatif puis une confirmation explicite."),
        ("États compréhensibles", "Chaque élément important affiche un état en mots simples : brouillon, en attente, accepté, refusé, actif, expiré, annulé, remboursé ou terminé."),
        ("Reprise après interruption", "Une candidature, un paiement non finalisé ou un formulaire long permet de comprendre ce qui a été conservé et comment reprendre."),
        ("Absence de contenu", "Une liste vide explique pourquoi elle est vide et propose une action utile : découvrir, créer, rechercher, modifier un filtre ou demander de l’aide."),
        ("Erreur et nouvel essai", "Un échec explique ce qui n’a pas abouti, conserve les informations sûres et propose de réessayer sans provoquer de doublon."),
        ("Accessibilité", "Les textes sont lisibles, les contrastes suffisants, les actions identifiables autrement que par la couleur et les parcours utilisables au clavier ou avec une aide de lecture."),
        ("Respect de la vie privée", "Les informations personnelles ne sont affichées que dans le contexte nécessaire. Les préférences, blocages, exports et demandes de suppression restent accessibles."),
        ("Assistance contextualisée", "Depuis un achat, un billet, une candidature ou un abonnement, la demande d’aide reprend le contexte utile afin d’éviter à la personne de tout réexpliquer."),
        ("Information après changement", "Lorsqu’une date, un prix, une attribution, une décision ou un état change, les personnes concernées voient la nouvelle information et, si nécessaire, reçoivent une alerte."),
    ]
    add_table(doc, ["Règle", "Comportement fonctionnel attendu"], transversal, [2700, 6660])

    doc.add_heading("4.1 Parcours de montée en rôle", level=2)
    add_table(doc, ["Départ", "Parcours", "Arrivée"], [
        ("Utilisateur", "Choisit Devenir organisateur → prépare et envoie sa candidature → suit l’examen → complète son profil.", "Utilisateur + Organisateur"),
        ("Utilisateur", "Choisit Devenir prestataire → prépare et envoie sa candidature → suit l’examen → publie son profil et ses services.", "Utilisateur + Prestataire"),
        ("Multi-rôle", "Ouvre le sélecteur de rôle → choisit l’espace souhaité → retrouve les menus et informations correspondants.", "Rôle actif changé"),
    ], [2100, 5160, 2100])

    doc.add_heading("4.2 Cycle de vie des éléments principaux", level=2)
    add_table(doc, ["Élément", "Cycle fonctionnel"], [
        ("Candidature", "Brouillon → envoyée → en examen → complément éventuel → acceptée ou refusée."),
        ("Événement", "Brouillon → prêt → publié → modifié, reporté ou annulé → terminé → bilan."),
        ("Billet", "Disponible → sélectionné → réservé ou payé → attribué éventuellement → contrôlé, revendu, rendu ou remboursé."),
        ("Invitation", "Envoyée → en attente → acceptée, refusée ou annulée."),
        ("Versement", "Revenus en attente → disponibles → demande envoyée → traitement → reçu ou à corriger."),
        ("Prestation", "Brouillon → publiée → modifiée ou masquée → republiée ou archivée."),
        ("Abonnement", "Choisi → paiement en cours → actif → renouvelé, fin programmée ou expiré."),
        ("Signalement", "Envoyé → reçu pour examen → décision et mesures appropriées lorsque nécessaires."),
    ], [2450, 6910])

    add_page_break(doc)
    doc.add_heading("5. Index fonctionnel", level=1)
    add_para(doc, "Cet index permet de retrouver rapidement le code de chaque fonctionnalité détaillée dans les parties précédentes.")
    for title, prefix, sections in [
        ("5.1 Utilisateur", "U", USER_SECTIONS),
        ("5.2 Organisateur", "O", ORGANIZER_SECTIONS),
        ("5.3 Prestataire", "P", PROVIDER_SECTIONS),
    ]:
        doc.add_heading(title, level=2)
        rows = []
        idx = 1
        for section_title, items in sections:
            for item in items:
                rows.append((f"{prefix}-{idx:03d}", item["title"], section_title.split(". ", 1)[-1]))
                idx += 1
        add_table(doc, ["Code", "Fonctionnalité", "Domaine"], rows, [1300, 4060, 4000])

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(1.25)
    section.bottom_margin = Inches(1.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    style_run(p.add_run("LIVEINBLACK"), size=14, bold=True, color="C83756")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("Un écosystème, trois rôles, des parcours explicites."), size=21, bold=True, color="3B1B23")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("Document fonctionnel de référence"), size=12, color="7E5E66")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
